#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FindJob V2.0 产品需求文档 (PRD) 生成器
完整版：13章 + 附录
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# 全局样式配置
# ============================================================

def setup_styles(doc):
    """配置文档基础样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    pf = style.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    h1.font.name = 'Microsoft YaHei'
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    h2.font.name = 'Microsoft YaHei'
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    h3.font.name = 'Microsoft YaHei'
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.size = Pt(11.5)
    h4.font.bold = True
    h4.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
    h4.font.name = 'Microsoft YaHei'
    h4.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)


def add_header_footer(doc):
    """添加页眉页脚"""
    for section in doc.sections:
        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = "FindJob PRD"
        hp.style.font.size = Pt(9)
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 页脚 - 页码
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar2)


# ============================================================
# 辅助函数
# ============================================================

def p(doc, text, bold=False, size=None, color=None, align=None, indent=None, space_after=None):
    """快速添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        para.alignment = align
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


def bullet(doc, text, level=0, bold_prefix=None):
    """添加列表项"""
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = para.add_run(bold_prefix)
        run_b.bold = True
        para.add_run(text)
    else:
        # Clear any existing runs and add text
        for run in para.runs:
            run.text = ''
        para.add_run(text)
    if level > 0:
        para.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return para


def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)
            for paragraph in cells[ci].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    # 列宽
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table


def status_tag(text, is_done=True):
    """生成状态标记文本"""
    return f"[已实现] {text}" if is_done else f"[规划中] {text}"


def add_feature_block(doc, fid, name, priority, status_done, desc, user_story,
                      preconditions, main_flow, alt_flows, exception_handling,
                      output_spec, prompt_example=None, acceptance=None, ui_desc=None):
    """添加完整的功能需求块（CareerAgent级别粒度）"""
    tag = "已实现" if status_done else "规划中"
    doc.add_heading(f'[{fid}] {name}  [{tag}]', level=3)

    # 基本信息表
    add_table(doc,
        ['属性', '内容'],
        [
            ['功能编号', fid],
            ['功能名称', name],
            ['优先级', priority],
            ['实现状态', '已实现' if status_done else '规划中'],
            ['描述', desc],
        ],
        col_widths=[3.5, 13]
    )

    # 用户故事
    doc.add_heading('用户故事', level=4)
    p(doc, user_story)

    # 前置条件
    doc.add_heading('前置条件', level=4)
    for item in preconditions:
        bullet(doc, item)

    # 主流程
    doc.add_heading('主流程', level=4)
    for i, step in enumerate(main_flow, 1):
        p(doc, f'{i}. {step}')

    # 替代流程
    if alt_flows:
        doc.add_heading('替代/分支流程', level=4)
        for af in alt_flows:
            bullet(doc, af)

    # 异常处理
    doc.add_heading('异常处理', level=4)
    for exc in exception_handling:
        bullet(doc, exc)

    # 输出规格
    doc.add_heading('输出规格', level=4)
    p(doc, output_spec)

    # Prompt 示例
    if prompt_example:
        doc.add_heading('Prompt 示例', level=4)
        para = doc.add_paragraph()
        run = para.add_run(prompt_example)
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        para.paragraph_format.left_indent = Cm(0.5)

    # 页面/交互描述
    if ui_desc:
        doc.add_heading('页面交互说明', level=4)
        p(doc, ui_desc)

    # 验收标准
    if acceptance:
        doc.add_heading('验收标准', level=4)
        for a in acceptance:
            bullet(doc, a)


# ============================================================
# 封面
# ============================================================

def add_cover(doc):
    """添加封面"""
    # 多个空行推到页面中部
    for _ in range(6):
        doc.add_paragraph()

    p(doc, 'FindJob V2.0', bold=True, size=36, align=WD_ALIGN_PARAGRAPH.CENTER,
      color=(0x1a, 0x1a, 0x1a))
    p(doc, '产品需求文档 (PRD)', bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER,
      color=(0x33, 0x33, 0x33))

    doc.add_paragraph()
    p(doc, 'AI 驱动的全流程求职操作系统 & 职业成长平台', size=14,
      align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x66, 0x66, 0x66))

    for _ in range(4):
        doc.add_paragraph()

    info_items = [
        ('文档版本', 'V2.0'),
        ('产品名称', 'FindJob'),
        ('文档类型', '产品需求文档 (PRD)'),
        ('编写日期', '2026-03-10'),
        ('文档状态', '评审中'),
        ('密级', '内部公开'),
    ]
    table = doc.add_table(rows=len(info_items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(info_items):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(11)
        table.rows[i].cells[0].width = Cm(5)
        table.rows[i].cells[1].width = Cm(8)

    doc.add_page_break()


# ============================================================
# 修订记录 + 目录
# ============================================================

def add_revision_and_toc(doc):
    """修订历史与目录"""
    doc.add_heading('修订记录', level=1)
    add_table(doc,
        ['版本', '日期', '修订人', '修订内容'],
        [
            ['V0.5', '2026-01-12', 'FindJob Team', '初版MVP PRD（JD解析+简历匹配）'],
            ['V0.7', '2026-02-01', 'FindJob Team', '新增面试准备、简历优化模块'],
            ['V1.0', '2026-02-15', 'FindJob Team', '完成Phase 1-2后端API，新增投递跟踪'],
            ['V1.5', '2026-03-01', 'FindJob Team', '高保真Figma设计完成，7大模块80+组件'],
            ['V2.0', '2026-03-10', 'FindJob Team', '全面升级：13章完整PRD，新增商业化路线图、竞品深度分析、Agent矩阵、Prompt工程体系'],
        ],
        col_widths=[2, 3, 3, 8.5]
    )

    doc.add_heading('目录', level=1)
    p(doc, '（在Word中请右键此处 → 更新域 → 更新整个目录）')
    # 插入TOC域
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._r.append(fldChar2)
    run2 = paragraph.add_run('【请在Word中右键更新目录】')
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2._r.append(fldChar3)
    doc.add_page_break()


# ============================================================
# 第1章 产品综述
# ============================================================

def add_chapter_1(doc):
    doc.add_heading('第一章 产品综述', level=1)

    # 1.1
    doc.add_heading('1.1 产品背景与市场机会', level=2)
    p(doc, '2025-2026年，全球就业市场正经历深刻变革。AI技术的快速渗透带来三大结构性变化：')
    bullet(doc, '信息不对称加剧：求职者平均每次求职需浏览200+岗位，但仅有不到5%的岗位真正匹配其能力画像。招聘方使用ATS（Applicant Tracking System）自动筛选简历，通过率不足25%。')
    bullet(doc, 'AI成本拐点到来：大语言模型（LLM）推理成本在2025年下降超过80%，使得"AI for Everyone"成为可能。单次API调用成本从2024年的$0.06降至2026年的$0.003。')
    bullet(doc, '求职效率瓶颈：传统求职流程（搜索→投递→等待→面试→决策）平均耗时3-6个月，其中60%时间浪费在低效信息处理上。')
    doc.add_paragraph()
    p(doc, 'FindJob V2.0 正是在这一背景下诞生，定位为 AI 驱动的全流程求职操作系统与职业成长平台，旨在用AI重构求职全链路，将"搜索式求职"升级为"AI辅助的精准求职"。')

    # 1.2 市场规模
    doc.add_heading('1.2 市场规模分析', level=2)
    add_table(doc,
        ['市场层级', '规模（2026E）', '增长率', '说明'],
        [
            ['TAM（总可用市场）', '¥8,500亿', '12%', '全球在线招聘市场'],
            ['SAM（可服务市场）', '¥420亿', '25%', '中国AI求职工具市场'],
            ['SOM（可获得市场）', '¥8.4亿', '40%', '目标用户群（中高端求职者AI工具付费意愿）'],
        ],
        col_widths=[4, 3, 2, 7.5]
    )
    p(doc, '核心驱动力：(1) 中国每年约900万高校毕业生+数千万社招求职者；(2) AI工具付费渗透率从2024年的3%预计增至2026年的15%；(3) 简历优化和面试辅导是付费意愿最高的两个场景。')

    # 1.3 竞品分析
    doc.add_heading('1.3 竞品深度分析', level=2)

    doc.add_heading('1.3.1 国际竞品', level=3)

    # Jobscan
    doc.add_heading('Jobscan', level=4)
    p(doc, '定位：ATS简历优化工具领军者')
    add_table(doc,
        ['维度', '详情'],
        [
            ['核心功能', 'ATS关键词匹配评分、简历扫描、LinkedIn优化、求职信生成'],
            ['技术特点', '基于NLP关键词提取+TF-IDF匹配算法，非LLM驱动'],
            ['定价', '免费版5次/月，Pro $49.95/月，Premium $89.95/月'],
            ['用户规模', '全球150万+用户，以北美市场为主'],
            ['优势', '品牌知名度高、ATS数据库全面（覆盖100+ATS系统）、简单易用'],
            ['劣势', '仅覆盖简历优化单环节、不支持中文、无面试辅导、无AI对话能力、匹配算法基于关键词非语义理解'],
        ],
        col_widths=[3, 13.5]
    )

    # Huntr
    doc.add_heading('Huntr', level=4)
    p(doc, '定位：求职流程管理工具（Job Tracker）')
    add_table(doc,
        ['维度', '详情'],
        [
            ['核心功能', '求职看板（Kanban）、简历生成器、求职信AI生成、面试准备笔记、联系人管理'],
            ['技术特点', '项目管理思维的求职工具，AI功能为辅助（GPT集成）'],
            ['定价', '免费版基础功能，Pro $40/月'],
            ['用户规模', '约50万用户'],
            ['优势', '看板式管理直观、浏览器插件一键保存岗位、多平台同步'],
            ['劣势', 'AI能力浅层（仅文本生成）、无简历匹配分析、无面试模拟、不支持中文市场'],
        ],
        col_widths=[3, 13.5]
    )

    # Kickresume
    doc.add_heading('Kickresume', level=4)
    p(doc, '定位：AI简历/求职信生成器')
    add_table(doc,
        ['维度', '详情'],
        [
            ['核心功能', 'AI简历生成、30+模板、求职信生成、个人网站生成'],
            ['技术特点', '基于GPT-4的文本生成，模板引擎渲染'],
            ['定价', '免费基础版，Premium $19/月，Ultimate $49/月'],
            ['优势', '模板丰富美观、AI生成速度快、支持多语言'],
            ['劣势', '仅覆盖简历制作、无岗位匹配分析、无面试准备、无求职流程管理'],
        ],
        col_widths=[3, 13.5]
    )

    # Teal
    doc.add_heading('Teal', level=4)
    p(doc, '定位：AI求职助手平台')
    add_table(doc,
        ['维度', '详情'],
        [
            ['核心功能', 'AI简历定制、求职追踪、技能匹配、职业规划'],
            ['技术特点', 'LLM+传统NLP混合方案，浏览器插件数据采集'],
            ['定价', '免费版，Teal+ $29/月'],
            ['优势', '功能较全面（求职追踪+简历+匹配）、Chrome插件体验好'],
            ['劣势', '面试准备模块弱、无中文支持、AI深度不足（无多Agent协作）'],
        ],
        col_widths=[3, 13.5]
    )

    # Rezi
    doc.add_heading('Rezi', level=4)
    p(doc, '定位：AI简历优化专家')
    add_table(doc,
        ['维度', '详情'],
        [
            ['核心功能', 'AI简历撰写、ATS检测、关键词优化、即时评分'],
            ['定价', '免费基础版，Pro $29/月，Lifetime $149'],
            ['优势', 'ATS优化专业度高、终身制定价有吸引力'],
            ['劣势', '功能单一（仅简历）、无求职管理、无面试辅导'],
        ],
        col_widths=[3, 13.5]
    )

    doc.add_heading('1.3.2 国内竞品', level=3)

    # 超级简历
    doc.add_heading('超级简历（WonderCV）', level=4)
    p(doc, '定位：中国最大在线简历平台')
    add_table(doc,
        ['维度', '详情'],
        [
            ['核心功能', '在线简历编辑器、模板库、AI润色、简历评分、求职信'],
            ['技术特点', '模板引擎+基础NLP，近期接入LLM能力'],
            ['定价', '免费基础版，会员 ¥98/季、¥198/年'],
            ['用户规模', '2000万+注册用户'],
            ['优势', '中文市场占有率第一、模板生态丰富（200+）、品牌认知度高'],
            ['劣势', 'AI能力停留在润色层面、无岗位深度匹配、无面试辅导、无求职管理、商业模式依赖模板和打印付费'],
        ],
        col_widths=[3, 13.5]
    )

    # 牛客
    doc.add_heading('牛客网', level=4)
    add_table(doc,
        ['维度', '详情'],
        [
            ['定位', '技术求职社区+面试题库'],
            ['核心功能', '在线编程题库、面经分享、模拟面试、企业内推'],
            ['优势', '技术类求职者首选平台、真实面经数据丰富'],
            ['劣势', '偏技术岗位、无简历优化、面试辅导非AI驱动、不覆盖求职全链路'],
        ],
        col_widths=[3, 13.5]
    )

    # 知页简历
    doc.add_heading('知页简历', level=4)
    add_table(doc,
        ['维度', '详情'],
        [
            ['定位', '简历设计工具'],
            ['核心功能', '可视化简历编辑、设计师模板、PDF导出'],
            ['优势', '设计感强、操作流畅'],
            ['劣势', '纯工具类、无AI能力、无匹配分析'],
        ],
        col_widths=[3, 13.5]
    )

    # 脉脉
    doc.add_heading('脉脉', level=4)
    add_table(doc,
        ['维度', '详情'],
        [
            ['定位', '职场社交+招聘平台'],
            ['核心功能', '职场社交、公司评价、招聘信息、薪资查询'],
            ['优势', '真实职场信息、公司内幕、人脉网络'],
            ['劣势', '不是求职工具、无AI辅助、无简历/面试功能'],
        ],
        col_widths=[3, 13.5]
    )

    # 竞品矩阵
    doc.add_heading('1.3.3 竞品能力矩阵对比', level=3)
    add_table(doc,
        ['能力维度', 'FindJob V2.0', 'Jobscan', 'Huntr', 'Teal', '超级简历', '牛客'],
        [
            ['JD深度解析', '★★★★★ 多Agent', '★★★ 关键词', '★ 无', '★★★ 基础', '★ 无', '★ 无'],
            ['简历智能匹配', '★★★★★ 语义+雷达', '★★★★ ATS关键词', '★ 无', '★★★ 基础', '★★ 评分', '★ 无'],
            ['定向简历生成', '★★★★★ 一键生成', '★★ 建议', '★★ AI辅助', '★★★ AI', '★★ 润色', '★ 无'],
            ['面试全链路', '★★★★★ 题库+模拟+复盘', '★ 无', '★★ 笔记', '★ 无', '★ 无', '★★★★ 题库'],
            ['Offer决策', '★★★★ 对比+AI建议', '★ 无', '★★ 追踪', '★★ 追踪', '★ 无', '★ 无'],
            ['职业成长', '★★★★ AI伴侣+Skills', '★ 无', '★ 无', '★★ 规划', '★ 无', '★★ 社区'],
            ['数据驾驶舱', '★★★★ 6KPI+漏斗', '★ 无', '★★ 统计', '★★ 统计', '★ 无', '★ 无'],
            ['Multi-Agent', '★★★★★ 15+Agent DAG', '★ 无', '★ 无', '★★ 单Agent', '★ 无', '★ 无'],
            ['中文支持', '★★★★★ 原生', '★ 无', '★ 无', '★ 无', '★★★★★ 原生', '★★★★★ 原生'],
        ],
        col_widths=[3, 2.5, 2.5, 2, 2, 2.5, 2]
    )

    # 1.3.4 竞争定位
    doc.add_heading('1.3.4 差异化竞争壁垒', level=3)
    p(doc, 'FindJob V2.0 的核心差异化优势：')
    bullet(doc, '全链路覆盖：业内唯一覆盖"JD解析→简历匹配→定向简历→面试准备→模拟面试→面试复盘→Offer决策→职业成长"全流程的AI产品。竞品最多覆盖2-3个环节。')
    bullet(doc, 'Multi-Agent协作：15+专业Agent通过DAG引擎协作，而非简单的单次LLM调用。每个Agent有独立的Prompt工程、输入/输出规范和质量监控。')
    bullet(doc, '中文原生+本地化：针对中国求职市场设计（支持BOSS直聘/拉勾/猎聘等平台JD格式、中文简历STAR法则、国内公司面试风格）。')
    bullet(doc, '数据闭环驱动：从用户行为到AI质量到商业指标的完整数据体系，支持持续优化。')
    bullet(doc, '成本优势：基于阿里云百炼（Qwen系列）的模型选型，单次任务成本 ¥2.8，远低于依赖GPT-4的竞品。')

    # 1.4 产品定位与核心价值
    doc.add_heading('1.4 产品定位与核心价值', level=2)
    p(doc, '产品Slogan：透视机会 · 定制武器 · 陪伴成长', bold=True, size=13)
    doc.add_paragraph()

    add_table(doc,
        ['价值维度', '核心能力', '用户价值'],
        [
            ['透视', 'JD深度解析 + 公司分析 + 匹配度诊断', '让用户"看懂"每一个岗位的真实需求'],
            ['定制', '定向简历生成 + 面试个性化准备', '为每个目标岗位打造专属武器'],
            ['陪伴', 'AI成长伴侣 + 长期记忆 + Skills自动化', '不是工具，是陪伴用户成长的职业伙伴'],
        ],
        col_widths=[2.5, 6.5, 7.5]
    )

    # 1.5 目标用户画像
    doc.add_heading('1.5 目标用户画像', level=2)

    doc.add_heading('P0 核心用户：定向求职者', level=3)
    add_table(doc,
        ['属性', '描述'],
        [
            ['画像', '25-35岁，1-8年工作经验，明确目标岗位和公司'],
            ['痛点', '简历千篇一律无法突出优势；面试准备没有方向；Offer选择缺乏数据支撑'],
            ['场景', '每周投递5-15个岗位，同时准备2-3个面试'],
            ['付费意愿', '高（月均愿付 ¥50-150），因为求职结果直接影响收入'],
            ['期望', '提高投递回复率、面试通过率、拿到更好的Offer'],
        ],
        col_widths=[3, 13.5]
    )

    doc.add_heading('P1 重要用户：迷茫探索者', level=3)
    add_table(doc,
        ['属性', '描述'],
        [
            ['画像', '应届生或转行者，不确定方向'],
            ['痛点', '不知道适合什么岗位；简历没有亮点；面试缺乏经验'],
            ['场景', '广撒网投递，需要AI帮助梳理方向'],
            ['付费意愿', '中等（月均 ¥30-80）'],
        ],
        col_widths=[3, 13.5]
    )

    doc.add_heading('P2 潜在用户：长期维护者', level=3)
    add_table(doc,
        ['属性', '描述'],
        [
            ['画像', '在职但关注市场的资深人士'],
            ['痛点', '需要持续更新简历；偶尔面试需快速准备'],
            ['场景', '每月浏览1-2次，被猎头联系时集中使用'],
            ['付费意愿', '中低，但LTV高（持续数年）'],
        ],
        col_widths=[3, 13.5]
    )

    # 1.6 产品范围
    doc.add_heading('1.6 V2.0 产品范围', level=2)
    p(doc, 'FindJob V2.0 包含以下核心内容：')
    add_table(doc,
        ['维度', '数量', '说明'],
        [
            ['功能模块', '7 大模块', '机会工作台、资产中心、面试工作台、决策中心、成长中心、首页、数据驾驶舱'],
            ['功能点', '27 个 (F1-F27)', 'P0: 15个 | P1: 8个 | P2: 4个'],
            ['AI Agent', '15+', '覆盖JD解析、简历、匹配、面试、成长等'],
            ['React 组件', '80+', '包含67个ShadCN UI基础组件'],
            ['页面', '15+', '7大模块页面 + 首页 + 驾驶舱'],
            ['API 接口', '90+', '按RESTful规范设计'],
            ['数据库表', '30+', 'PostgreSQL 14+'],
            ['Prompt 模板', '12+', '结构化Prompt工程体系'],
            ['PRD 文档', '14份', '各模块高保真设计文档'],
        ],
        col_widths=[3, 3, 10.5]
    )

    # 1.7 流程图 Prompt
    doc.add_heading('1.7 业务流程图制作指引', level=2)
    p(doc, '以下Prompt供设计团队使用AI工具（如Mermaid/Draw.io/Figma）生成业务流程图：')

    prompts = [
        ('全生命周期流程图',
         '请绘制一个求职全生命周期流程图，包含以下阶段：\n'
         '1. 岗位发现阶段：用户输入JD/URL → JD预处理（文本清洗/图片OCR）→ JD结构化提取 → A维度分析（技术栈/产品类型/业务领域/团队阶段）→ B维度分析（行业/技术/经验/能力）→ 公司分析 → 保存到岗位池\n'
         '2. 武器准备阶段：选择基础简历 → 匹配度评估（5维雷达图）→ 定向简历生成（关键词注入+差距弥补+亮点强化）→ 版本管理 → 导出PDF\n'
         '3. 面试备战阶段：AI生成面试题 → 题库练习+AI评分 → 模拟面试 → 面试记录 → AI复盘（多维评分+优化建议）\n'
         '4. 决策阶段：录入Offer → 多维对比（薪资/发展/团队/文化/城市）→ AI建议 → 谈薪辅助\n'
         '5. 成长阶段：AI成长伴侣对话 → 周计划/周复盘 → Skills自动化 → 长期记忆积累\n'
         '风格：白色背景，灰色线条，每个阶段用不同灰度区分，节点使用圆角矩形，AI节点添加星号标记。'),

        ('Multi-Agent DAG调度流程图',
         '请绘制一个DAG（有向无环图）调度流程图，展示以下Agent的依赖关系：\n'
         '层级0（无依赖，可并行）：JD预处理Agent、简历解析Agent\n'
         '层级1（依赖层级0）：JD结构化Agent（依赖JD预处理）\n'
         '层级2（依赖层级1）：A维度分析Agent、B维度分析Agent（均依赖JD结构化，可并行）\n'
         '层级3（依赖层级2）：公司分析Agent（依赖A分析）、匹配评估Agent（依赖A+B分析+简历解析）\n'
         '层级4（依赖层级3）：简历优化Agent（依赖匹配评估）、面试准备Agent（依赖匹配评估+公司分析）\n'
         '标注每个Agent的模型选择（如qwen-max/qwen-plus/qwen-turbo）和预估耗时。'),

        ('数据闭环漏斗图',
         '请绘制一个6层漏斗图，展示用户求职闭环：\n'
         '第1层：进入首页（12,840用户，100%）\n'
         '第2层：完成岗位解析（8,526用户，66.4%）\n'
         '第3层：生成定向简历（5,116用户，39.8%）\n'
         '第4层：开始面试训练（3,719用户，29.0%）\n'
         '第5层：建立投递/面试记录（2,568用户，20.0%）\n'
         '第6层：进入Offer/决策阶段（1,283用户，10.0%）\n'
         '在漏斗右侧标注每层转化率和流失原因，底部显示北极星指标：求职闭环率 = Offer阶段用户/首页用户 = 10%，目标32%。'),
    ]

    for title, prompt_text in prompts:
        doc.add_heading(title, level=3)
        para = doc.add_paragraph()
        run = para.add_run(prompt_text)
        run.font.size = Pt(9.5)
        run.font.name = 'Consolas'
        para.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()


# ============================================================
# 第2章 系统架构
# ============================================================

def add_chapter_2(doc):
    doc.add_heading('第二章 业务流程与系统架构', level=1)

    # 2.1
    doc.add_heading('2.1 求职全生命周期流程', level=2)
    p(doc, 'FindJob V2.0 覆盖求职者从"发现机会"到"职业成长"的完整生命周期，共分为五个核心阶段：')
    doc.add_paragraph()

    add_table(doc,
        ['阶段', '核心动作', '涉及模块', 'AI介入点'],
        [
            ['1. 岗位发现', 'JD采集 → 深度解析 → 公司分析 → 匹配诊断', '机会工作台', 'JD预处理Agent、结构化Agent、A/B维度分析Agent、公司分析Agent'],
            ['2. 武器准备', '简历管理 → 匹配评估 → 定向简历生成 → 版本管理', '资产中心', '简历解析Agent、匹配评估Agent、简历优化Agent、版本生成Agent'],
            ['3. 面试备战', '题库练习 → 面试辅导 → 模拟面试 → 实战记录 → AI复盘', '面试工作台', '面试准备Agent、面试教练Agent、复盘Agent'],
            ['4. 决策阶段', 'Offer录入 → 多维对比 → 谈薪辅助 → AI选择建议', '决策中心', 'Offer分析Agent'],
            ['5. 持续成长', 'AI对话 → 周计划/复盘 → Skills自动化 → 长期记忆', '成长中心', '成长伴侣Agent、周复盘Agent'],
        ],
        col_widths=[2.5, 4.5, 2.5, 7]
    )

    p(doc, '（请参照第1.7节流程图Prompt，由设计团队生成对应的高保真流程图。）')

    # 2.2 系统架构总览
    doc.add_heading('2.2 系统架构总览', level=2)
    p(doc, 'FindJob V2.0 采用前后端分离 + Serverless + Multi-Agent 的现代化架构：')

    add_table(doc,
        ['层级', '技术选型', '说明'],
        [
            ['前端展示层', 'React 18 + TypeScript + Tailwind CSS v4 + ShadCN UI', '80+组件，React Router v7 路由管理'],
            ['应用框架层', 'Hono（轻量级Web框架）', '统一处理API路由、SSR页面、中间件'],
            ['AI服务层', 'Multi-Agent DAG引擎', '15+Agent，支持并行执行、重试、成本监控'],
            ['LLM接入层', '阿里云百炼（Dashscope）+ VectorEngine', '多模型调度：Qwen-Max/Plus/Turbo/VL'],
            ['数据层', 'PostgreSQL 14+ / Redis', '30+表，含视图、触发器、索引优化'],
            ['部署层', 'Cloudflare Pages + Workers', '全球CDN加速，Serverless无服务器架构'],
            ['构建工具', 'Vite 6 + Wrangler', '双构建：SPA客户端 + Worker服务端'],
        ],
        col_widths=[3, 6, 7.5]
    )

    # 2.3 前端架构
    doc.add_heading('2.3 前端架构详解', level=2)

    doc.add_heading('2.3.1 技术栈', level=3)
    add_table(doc,
        ['技术', '版本', '用途'],
        [
            ['React', '18.3.1', '核心UI框架'],
            ['TypeScript', '5.x', '类型安全'],
            ['Tailwind CSS', 'v4.1.12', '原子化CSS，含自定义设计系统变量'],
            ['ShadCN UI', '最新', '67个基础组件（Radix UI + CVA）'],
            ['React Router', 'v7 (createBrowserRouter)', 'SPA路由管理'],
            ['Recharts', '最新', '图表可视化（驾驶舱用）'],
            ['Lucide React', '最新', '图标库'],
            ['React Hook Form', '最新', '表单管理'],
            ['React DnD', '最新', '拖拽交互'],
        ],
        col_widths=[4, 3, 9.5]
    )

    doc.add_heading('2.3.2 设计系统', level=3)
    p(doc, 'FindJob设计系统定义在theme.css中，统一管理颜色、圆角、阴影、间距和排版：')
    add_table(doc,
        ['设计Token', '值', '说明'],
        [
            ['--background', '#fafaf9', '页面背景（米白色）'],
            ['--foreground', '#1a1a1a', '主文字颜色'],
            ['--card', '#ffffff', '卡片背景（纯白）'],
            ['--radius-card-lg', '28px', '大卡片圆角'],
            ['--radius-card-md', '20px', '中卡片圆角'],
            ['--radius-button', '14px', '按钮圆角'],
            ['--radius-pill', '999px', '胶囊形圆角'],
            ['页面边距', '28px', '页面内容与边缘间距'],
            ['卡片间距', '20-24px', '卡片之间的间距'],
        ],
        col_widths=[4, 3, 9.5]
    )

    doc.add_heading('2.3.3 页面路由结构', level=3)
    add_table(doc,
        ['路径', '组件', '说明'],
        [
            ['/', 'HomePage', '首页（今日工作台）'],
            ['/opportunities', 'OpportunitiesPage', '机会工作台'],
            ['/assets', 'AssetsPage', '资产中心'],
            ['/interviews', 'InterviewsPage', '面试工作台'],
            ['/decisions', 'DecisionsPage', '决策中心'],
            ['/growth', 'GrowthPage', '成长中心'],
            ['/dashboard', 'DashboardPage', '数据驾驶舱'],
        ],
        col_widths=[4, 5, 7.5]
    )

    # 2.4 后端架构
    doc.add_heading('2.4 后端架构详解', level=2)

    doc.add_heading('2.4.1 Hono框架 + Cloudflare Workers', level=3)
    p(doc, 'FindJob后端基于Hono框架运行在Cloudflare Workers上，采用Serverless架构：')
    bullet(doc, '路由层：Hono处理所有HTTP请求，区分API路由（/api/*）和页面路由')
    bullet(doc, 'API路由模块：job、resume、interview、optimize、metrics、questions、applications、feishu、market、chat，共10个路由模块')
    bullet(doc, '中间件：CORS跨域、请求日志、Feishu配置初始化')
    bullet(doc, '环境变量：PYTHON_SERVICE_URL（PDF解析服务）、FEISHU_*（飞书集成）')

    doc.add_heading('2.4.2 LLM配置中心', level=3)
    p(doc, 'FindJob的LLM配置采用集中管理的Agent配置模式：')
    add_table(doc,
        ['配置项', '说明', '当前值'],
        [
            ['API提供商', '支持dashscope和vectorengine双通道', 'dashscope（主）+ vectorengine（备）'],
            ['默认模型', 'Agent默认使用的模型', 'qwen-plus'],
            ['默认Temperature', '生成随机性', '0.7'],
            ['默认MaxTokens', '最大输出Token', '4096'],
            ['联网搜索', '部分Agent启用', 'chat-agent启用'],
        ],
        col_widths=[3.5, 6, 7]
    )

    # 2.5 DAG引擎
    doc.add_heading('2.5 Multi-Agent DAG 执行引擎', level=2)
    p(doc, 'FindJob的核心技术创新是DAG（有向无环图）Agent调度引擎，支持多Agent并行执行和依赖管理：')
    doc.add_paragraph()

    p(doc, '核心能力：', bold=True)
    bullet(doc, '依赖解析：自动分析Agent间的输入/输出依赖关系')
    bullet(doc, '并行执行：无依赖的Agent自动并行执行（如A维度和B维度分析）')
    bullet(doc, '重试机制：单个Agent失败后自动重试（最多3次）')
    bullet(doc, '状态追踪：实时追踪每个Agent的状态（pending/running/completed/failed）')
    bullet(doc, '成本监控：记录每个Agent的Token消耗和响应时间')
    bullet(doc, '实验管理：支持A/B测试不同的Prompt版本')
    doc.add_paragraph()

    p(doc, '典型DAG执行流程（JD解析场景）：', bold=True)
    add_table(doc,
        ['执行层级', 'Agent', '依赖', '模型', '并行'],
        [
            ['层级0', 'JD预处理Agent', '无', 'qwen-turbo / qwen-vl-max', '可与简历解析并行'],
            ['层级1', 'JD结构化Agent', 'JD预处理', 'qwen-plus', '串行'],
            ['层级2', 'A维度分析Agent', 'JD结构化', 'qwen-max', '与B维度并行'],
            ['层级2', 'B维度分析Agent', 'JD结构化', 'qwen-max', '与A维度并行'],
            ['层级3', '公司分析Agent', 'A维度分析', 'qwen-max', '与匹配评估并行'],
            ['层级3', '匹配评估Agent', 'A+B分析+简历解析', 'qwen-max', '与公司分析并行'],
            ['层级4', '简历优化Agent', '匹配评估', 'qwen-max', '与面试准备并行'],
            ['层级4', '面试准备Agent', '匹配评估+公司分析', 'qwen-plus', '与简历优化并行'],
        ],
        col_widths=[2.5, 4, 4, 3.5, 2.5]
    )


# ============================================================
# 第3章 机会工作台
# ============================================================

def add_chapter_3(doc):
    doc.add_heading('第三章 机会工作台', level=1)
    p(doc, '机会工作台是FindJob的入口模块，帮助用户"收集岗位、理解岗位、推进值得投入的机会"。')

    # F1
    add_feature_block(doc,
        fid='F1', name='JD采集与录入', priority='P0', status_done=True,
        desc='支持用户通过多种方式录入岗位JD，包括文本粘贴、URL导入和图片OCR识别。系统自动进行预处理（文本清洗或图片识别），为后续深度解析做准备。',
        user_story='作为一名求职者，我希望能快速将看到的岗位信息录入系统，无论是复制的文本、招聘网站链接，还是截图，系统都能正确识别和处理。',
        preconditions=[
            '用户已登录系统',
            '系统LLM服务可用（dashscope/vectorengine）',
        ],
        main_flow=[
            '用户进入机会工作台，点击"添加岗位"按钮',
            '选择录入方式：文本粘贴 / URL导入 / 图片上传',
            '文本模式：粘贴JD文本 → 系统调用JD预处理Agent（文本清洗）→ 去除冗余格式、修正乱码、统一结构',
            '图片模式：上传截图 → 系统调用JD预处理Agent（OCR模式，使用qwen-vl-max视觉模型）→ 识别图片中的文字 → 输出结构化文本',
            'URL模式：输入招聘页面URL → 系统爬取页面内容 → 提取JD文本 → 进入文本清洗流程',
            '预处理完成后，自动触发JD结构化提取（F2）',
        ],
        alt_flows=[
            '用户可跳过预处理直接手动填写岗位信息',
            '图片质量差时，系统提示"识别结果可能不完整，请检查并补充"',
            'URL爬取失败时，提示用户改用文本粘贴方式',
        ],
        exception_handling=[
            '异常E1：文本内容少于50字符 → 提示"内容过短，请确认是否为完整JD"',
            '异常E2：图片模型调用失败 → 自动降级到文本模式，提示用户手动输入',
            '异常E3：URL访问被拒绝（反爬） → 提示"无法自动获取，请手动复制JD内容"',
            '异常E4：LLM服务不可用 → 重试3次后提示"AI服务暂时不可用，请稍后重试"',
        ],
        output_spec='清洗后的JD纯文本，保留标题层级和列表格式。字段：cleanedText(string), sourceType(text/image/url), preprocessDuration(ms)。',
        prompt_example='''System Prompt (JD预处理-图片模式):
你是一个专业的 OCR 和文本提取专家。

## 核心任务
从图片中提取完整的岗位描述（JD）文本。

## 提取规则
1. 完整提取：识别图片中所有 JD 相关文字
2. 结构保持：保留原有的标题、分段、列表结构
3. 关键信息优先：岗位名称、公司名称、岗位职责、任职要求、加分项、薪资福利
4. 噪声过滤：忽略广告、水印、页眉页脚

## 输出格式
直接输出提取的纯文本，保持原有结构层次。''',
        acceptance=[
            '文本模式：输入合法JD文本后2秒内返回清洗结果',
            '图片模式：支持PNG/JPG/WEBP，文件大小≤10MB，5秒内返回OCR结果',
            'URL模式：支持主流招聘平台（BOSS直聘、拉勾、猎聘）URL解析',
            '清洗后文本保留原有的结构层级（标题、列表、段落）',
        ],
        ui_desc='机会工作台顶部工具栏右侧"添加岗位"按钮，点击弹出录入Modal，提供三种Tab：文本粘贴、URL导入、图片上传。底部显示处理状态和进度。'
    )

    # F2
    add_feature_block(doc,
        fid='F2', name='JD深度解析', priority='P0', status_done=True,
        desc='对预处理后的JD文本进行多层次AI分析：结构化提取（岗位名称、职责、要求等字段）→ A维度分析（技术栈、产品类型、业务领域、团队阶段）→ B维度分析（行业要求、技术要求、产品经验、核心能力）。',
        user_story='作为一名求职者，我希望系统不仅能提取JD的基本信息，更能帮我"读懂"岗位背后的隐性需求，比如这个岗位实际需要什么样的人。',
        preconditions=[
            'JD文本已完成预处理（F1）',
            'LLM服务可用',
        ],
        main_flow=[
            '系统接收预处理后的JD文本',
            '第一步：JD结构化Agent → 提取title/company/location/salary/responsibilities/requirements/preferred/others',
            '第二步（并行执行）：A维度分析Agent → 分析A1技术栈（关键词+密度）、A2产品类型、A3业务领域、A4团队阶段',
            '第三步（并行执行）：B维度分析Agent → 分析B1行业背景要求、B2技术背景要求（学历+技术深度分级）、B3产品经验要求、B4核心能力要求',
            '结构化数据存储到岗位记录的parsed_data字段（JSONB）',
            '分析结果存储到岗位记录的analysis字段（JSONB）',
            '在岗位详情页展示完整解析结果',
        ],
        alt_flows=[
            'JD信息不完整时，Agent标注"信息不足"而非编造数据',
            '解析结果支持用户手动编辑修正',
        ],
        exception_handling=[
            '异常E1：结构化提取失败 → 返回原始文本 + 提示"自动解析失败，请手动填写关键信息"',
            '异常E2：A/B维度分析超时（>15秒） → 先返回结构化结果，维度分析异步完成后更新',
            '异常E3：JSON解析失败 → Agent输出非法JSON → 重试1次，仍失败则记录日志并返回默认结构',
            '异常E4：模型幻觉 → 输出的技术栈包含JD中未提及的技术 → 通过Validator检查，标注可信度',
        ],
        output_spec='''结构化JD：{title, company, location, salary, responsibilities[], requirements[], preferred[], others}
A维度：{A1_tech_stack: {keywords[], density, summary}, A2_product_type: {type, reason}, A3_business_domain: {primary, secondary[], summary}, A4_team_stage: {stage, evidence[], summary}}
B维度：{B1_industry_requirement: {required, preferred, years, specific_industry, summary}, B2_tech_requirement: {education, tech_depth: {了解[], 熟悉[], 精通[]}, summary}, B3_product_experience: {product_types[], need_full_cycle, need_0to1, summary}, B4_capability_requirement: {capabilities: [{name, detail}], summary}}''',
        prompt_example='''System Prompt (A维度分析):
你是资深招聘分析专家，擅长分析岗位定位。

## 任务
对 JD 进行 A 维度分析（岗位定位），包含 4 个子维度：
- A1 技术栈分析：提取技术关键词，判断技术密度(高/中/低)
- A2 产品类型：判断ToB/ToC/平台/工具/数据/AI产品
- A3 业务领域：primary主领域 + secondary次领域
- A4 团队阶段：0-1搭建期/1-10成长期/10-100成熟期/维护优化期

## 分析原则
1. 基于 JD 原文分析，不编造信息
2. 技术词仅提取明确提及的
3. 信息不足时标注"JD未明确提及"

User Message:
请对以下岗位进行A维度分析：
## 岗位基本信息 - 岗位名称：{title} - 公司：{company}
## 岗位职责 1. {responsibility_1} ...
## 任职要求 1. {requirement_1} ...''',
        acceptance=[
            '结构化提取准确率≥90%（岗位名称、公司名称、职责列表）',
            'A维度分析在5秒内完成',
            'B维度分析在5秒内完成',
            '技术栈关键词提取覆盖JD中明确提及的技术词≥85%',
        ],
        ui_desc='岗位详情页右侧区域，通过Tab切换查看：概览、JD解析、能力模型、匹配诊断、公司分析、投递记录。解析结果以卡片+标签形式展示。'
    )

    # F3
    add_feature_block(doc,
        fid='F3', name='公司深度分析', priority='P0', status_done=True,
        desc='基于A维度分析结果和公开信息，对目标公司进行全方位分析：公司概况、AI应用场景、竞品格局、面试洞察。帮助求职者在面试前全面了解公司。',
        user_story='作为一名求职者，在面试前我想全面了解目标公司的业务、技术方向和面试风格，这样我能有针对性地准备。',
        preconditions=['A维度分析已完成', 'LLM联网搜索可用'],
        main_flow=[
            '系统基于A维度分析结果提取公司名和行业信息',
            '调用公司分析Agent（启用联网搜索获取公开信息）',
            '输出公司画像：行业、阶段、规模、主营产品、商业模式',
            '输出AI场景：当前AI应用、潜在AI机会、行业AI趋势、岗位AI工作重点',
            '输出竞品分析：直接竞品、竞争地位、差异化点',
            '输出面试洞察：公司文化、关键挑战、成长机会、面试建议',
        ],
        alt_flows=['公司为初创企业公开信息少 → Agent标注信息来源可信度'],
        exception_handling=[
            '异常E1：联网搜索失败 → 仅基于JD内容推断，标注"以下分析基于JD推断，建议自行核实"',
            '异常E2：公司名未识别 → 提示用户手动输入公司名',
        ],
        output_spec='''{company_profile: {name, industry, stage, scale, main_products[], business_model}, ai_scenarios: {current_ai_usage[], potential_ai_opportunities[], industry_ai_trends[], role_ai_focus}, competitor_analysis: {direct_competitors: [{name, description}], competitive_position, differentiation_points[]}, interview_insights: {company_culture, key_challenges[], growth_opportunities[], interview_tips[]}, summary}''',
        prompt_example='''System Prompt (公司分析Agent):
你是资深行业分析师和求职顾问。

## 任务
分析公司背景、AI应用场景和竞争格局，帮助求职者准备面试。

## 分析原则
1. 基于公开信息合理推断
2. AI场景分析要具体、可执行
3. 面试建议要针对岗位特点
4. 每个列表字段 2-4 项''',
        acceptance=[
            '分析结果在10秒内返回',
            '公司画像字段完整度≥80%',
            '面试建议具有针对性（与岗位相关，非通用建议）',
        ],
        ui_desc='岗位详情页 > 公司分析Tab。卡片式展示：顶部公司概况卡，下方分为AI场景、竞品格局、面试洞察三个区域。'
    )

    # F4
    add_feature_block(doc,
        fid='F4', name='匹配度诊断', priority='P0', status_done=True,
        desc='基于用户简历和JD的A/B维度分析结果，进行多维度匹配评估。生成匹配等级、匹配分数、5维雷达图（业务领域、行业背景、技术能力、产品经验、核心能力）、优势/差距分析和面试重点建议。',
        user_story='作为一名求职者，我想知道自己与目标岗位的匹配程度，了解我的优势在哪里、差距在哪里，以便有针对性地准备。',
        preconditions=['用户已上传并解析简历', 'JD已完成A/B维度分析'],
        main_flow=[
            '系统获取用户简历的结构化数据（基本信息、教育、工作经历、项目、技能、能力标签）',
            '系统获取目标岗位的结构化JD和A/B维度分析结果',
            '调用匹配评估Agent，输入简历+岗位数据',
            'Agent按规则计算匹配等级：非常匹配(A3一致+B1-B4全满足) / 比较匹配(A3一致+3项满足) / 匹配度还可以(A3相关+2项) / 不是很匹配 / 不匹配',
            '输出各维度匹配状态（✅匹配/⚠️部分/❌不匹配）、优势列表、差距列表、面试重点建议',
            '前端生成5维雷达图展示匹配度',
        ],
        alt_flows=[
            '用户无简历 → 提示先在资产中心上传简历',
            '简历信息不完整 → 标注"简历信息不完整，匹配结果仅供参考"',
        ],
        exception_handling=[
            '异常E1：匹配评估超时 → 返回基础匹配（仅关键词匹配），标注"深度分析正在进行"',
            '异常E2：简历与JD行业差距极大 → 仍给出分析，但在summary中说明"跨行业求职建议"',
        ],
        output_spec='''{match_level: "非常匹配"|"比较匹配"|"匹配度还可以"|"不是很匹配"|"不匹配", match_score: 0-100, dimension_match: {A3_business_domain: {status, evidence, detail}, B1_industry: {...}, B2_tech: {...}, B3_product: {...}, B4_capability: {...}}, strengths: string[], gaps: string[], interview_focus_suggestion: string}''',
        prompt_example='''匹配等级规则（严格遵循）：
| 等级 | 条件 |
| 非常匹配 | A3业务领域一致 + B1-B4全部满足(4项) |
| 比较匹配 | A3一致 + B1-B4满足3项 |
| 匹配度还可以 | A3相关 + B1-B4满足2项 |
| 不是很匹配 | A3不相关 或 B1-B4仅满足1项 |
| 不匹配 | B1-B4均不满足 |

评分标准：
- 85-100分：非常匹配
- 70-84分：比较匹配
- 55-69分：匹配度还可以
- 0-54分：匹配度较低''',
        acceptance=[
            '匹配评估在10秒内返回',
            '匹配等级与评分一致（如85分应为"非常匹配"）',
            '优势和差距各至少3条',
            '雷达图正确渲染5个维度',
        ],
        ui_desc='岗位详情页 > 匹配诊断Tab。顶部显示总分和匹配等级标签。左侧雷达图展示5维匹配度。右侧展示优势（绿色标签）和差距（橙色标签）。底部面试建议卡片。'
    )

    # F5
    add_feature_block(doc,
        fid='F5', name='定向简历生成', priority='P0', status_done=True,
        desc='基于匹配评估结果，一键为目标岗位生成定向优化的简历。核心策略：关键词注入（自然融入JD关键词）、差距弥补（调整表述弥补不足）、亮点强化（量化成果，STAR法则）。',
        user_story='作为一名求职者，我希望系统能根据每个目标岗位自动调整我的简历，突出相关经验，弥补不足，提高简历被ATS通过和HR关注的概率。',
        preconditions=['用户有至少一份基础简历', '匹配评估已完成'],
        main_flow=[
            '用户在岗位详情页点击"生成定向简历"按钮',
            '选择基础简历（默认使用主简历）',
            '可选：填写重点突出的能力方向',
            '系统调用简历优化Agent，输入：基础简历+JD分析+A/B维度+匹配结果',
            'Agent执行优化策略：(a)关键词注入 (b)差距弥补 (c)亮点强化 (d)表述优化（主动动词+量化数据）',
            '输出优化摘要、各section的original/optimized对比、修改说明、关键词覆盖率变化',
            '生成新简历版本，关联到目标岗位',
            '用户可预览对比、编辑、接受/拒绝各项修改',
        ],
        alt_flows=[
            '用户自定义修改建议 → Agent结合用户建议重点调整',
            '匹配度已经很高(>85分) → 仅做微调和表述优化',
        ],
        exception_handling=[
            '异常E1：简历内容过少无法优化 → 提示"简历内容不足，请先补充工作经历和项目经验"',
            '异常E2：优化后文本与原文差异过大 → 标注"以下修改较大，请仔细确认是否符合实际"',
            '异常E3：Agent输出格式异常 → 返回原始简历+错误提示，不覆盖原文',
        ],
        output_spec='''{optimization_summary: string, sections: {summary?: OptimizedSection, work_experience: [{company, position, original, optimized, changes[], matched_requirements[], keywords_added[]}], projects: [{name, original, optimized, changes[], matched_requirements[], keywords_added[]}], skills?: {original[], optimized[], added[], emphasized[], changes[]}}, optimization_effect: {keywords_coverage, gaps_addressed[], highlights_strengthened[], estimated_match_improvement}}''',
        prompt_example='''System Prompt (简历优化Agent):
你是资深简历优化专家。

## 优化策略
1. 关键词注入：自然融入JD关键词
2. 差距弥补：调整表述弥补不足
3. 亮点强化：量化成果，使用STAR法则
4. 表述优化：主动动词开头，结构清晰

## 优化原则
1. 保持真实性，不编造经历
2. 每个修改有明确理由
3. 优先优化硬性要求匹配点''',
        acceptance=[
            '定向简历在15秒内生成完成',
            '关键词覆盖率提升≥20%',
            '每个修改项都有对应的matched_requirements说明',
            '优化后文本保持自然流畅，无明显AI痕迹',
        ],
        ui_desc='岗位详情页顶部"生成定向简历"按钮。点击后弹出选择基础简历的Modal，确认后进入简历对比视图：左侧原文、右侧优化后，差异高亮标注。'
    )

    # F6
    add_feature_block(doc,
        fid='F6', name='岗位池管理', priority='P1', status_done=False,
        desc='管理所有已采集的岗位，支持搜索、筛选、排序、状态流转（saved→applied→interviewed→rejected）、标签管理和批量操作。',
        user_story='作为一名求职者，我希望能方便地管理所有采集的岗位，快速找到高匹配岗位，追踪每个岗位的投递状态。',
        preconditions=['用户已采集至少一个岗位'],
        main_flow=[
            '进入机会工作台，左侧显示岗位列表',
            '支持按关键词搜索（公司名/岗位名/技能标签）',
            '支持筛选：状态(saved/applied/interviewed)、匹配度范围、城市、薪资范围',
            '支持排序：创建时间、匹配度、公司名',
            '点击岗位卡片查看详情',
            '状态流转：saved → applied → interviewed → rejected/offered',
            '标签管理：添加/删除自定义标签',
        ],
        alt_flows=['岗位列表为空 → 显示引导页，提示添加第一个岗位'],
        exception_handling=[
            '异常E1：列表加载超时 → 显示缓存数据 + "加载中"提示',
            '异常E2：筛选结果为空 → 显示"未找到符合条件的岗位，请调整筛选条件"',
        ],
        output_spec='分页列表，每页20条：[{id, company, position, salaryRange, location, matchScore, status, tags[], createdAt}]',
        acceptance=[
            '列表加载时间<2秒',
            '支持至少5种筛选条件组合',
            '状态流转操作即时生效',
            '搜索支持模糊匹配',
        ],
        ui_desc='机会工作台左侧列表区域。每个岗位卡片显示：岗位名、公司名、城市、薪资、匹配度标签、状态标签。支持列表/卡片视图切换。'
    )


# ============================================================
# 第4章 资产中心
# ============================================================

def add_chapter_4(doc):
    doc.add_heading('第四章 资产中心', level=1)
    p(doc, '资产中心是用户的"求职武器库+版本系统+素材中心"，管理求职画像、简历库、项目素材和成就证据。核心理念：简历不是一份文档，而是持续迭代的资产。')

    # F7
    add_feature_block(doc,
        fid='F7', name='求职画像管理', priority='P0', status_done=True,
        desc='用户的基础信息中心，包含个人信息、职业目标、能力标签（行业/技术/产品/能力四维）。求职画像是所有AI分析的基础数据源。',
        user_story='作为一名求职者，我希望有一个地方集中管理我的职业信息，包括我的能力标签和求职目标，系统能基于这些信息提供个性化服务。',
        preconditions=['用户已注册并登录'],
        main_flow=[
            '用户进入资产中心，左侧导航选择"求职画像"',
            '填写/编辑基本信息：姓名、联系方式、目标岗位、期望薪资、期望城市',
            '管理能力标签：行业标签（如互联网/金融）、技术标签（如Python/SQL）、产品标签（如ToB/SaaS）、能力标签（如需求分析/数据驱动）',
            '上传简历后自动提取并补充能力标签',
            '求职画像数据作为匹配评估、简历优化的输入',
        ],
        alt_flows=['首次使用 → 引导用户完成画像设置（上传简历自动提取或手动填写）'],
        exception_handling=[
            '异常E1：简历解析提取标签失败 → 提示手动添加能力标签',
        ],
        output_spec='{basic_info: {name, phone, email, location, target_position, expected_salary}, ability_tags: {industry[], technology[], product[], capability[]}}',
        acceptance=[
            '能力标签支持自定义添加',
            '简历上传后自动提取标签准确率≥80%',
            '画像数据变更实时同步到匹配评估',
        ],
        ui_desc='资产中心左侧导航 > 求职画像。右侧编辑区域分为基本信息表单和能力标签管理区（四维标签以颜色区分的胶囊标签展示）。'
    )

    # F8
    add_feature_block(doc,
        fid='F8', name='简历库与版本管理', priority='P0', status_done=True,
        desc='支持多简历管理和版本控制。简历类型包括：基础简历（主简历）、通用优化版、定向简历（关联岗位）、手动编辑版。每个简历支持版本历史。',
        user_story='作为一名求职者，我希望能管理多份简历，每份简历对应不同岗位或版本，能随时回溯和对比历史版本。',
        preconditions=['用户已登录', '至少有一份基础简历'],
        main_flow=[
            '进入资产中心 > 简历库',
            '左侧树形导航展示简历分类：主简历、通用优化版、定向简历（按岗位分组）、手动编辑版',
            '点击简历查看详情：标题、状态标签、关联岗位、最后更新时间',
            '版本管理：每次编辑自动创建新版本、支持版本对比（diff视图）',
            '操作：编辑、另存为新版本、导出PDF/DOCX、删除',
            '定向简历自动关联到对应岗位（通过based_on_opportunity_id字段）',
        ],
        alt_flows=[
            '导入外部简历（PDF/DOCX上传） → 调用简历解析Agent → 自动提取结构化数据',
            '从定向简历生成（F5）自动创建新版本',
        ],
        exception_handling=[
            '异常E1：简历数量超限（最多20份） → 提示"已达上限，请删除旧简历后再创建"',
            '异常E2：PDF解析失败 → 调用备用PDF解析服务（PYTHON_SERVICE_URL）',
            '异常E3：导出PDF失败 → 提供DOCX格式作为备选',
        ],
        output_spec='简历列表：[{id, title, isMaster, version, basedOnOpportunityId, downloadCount, content: {basic, education, workExperience, projects, skills}, createdAt, updatedAt}]',
        prompt_example='''简历解析Agent System Prompt:
你是专业的简历信息提取专家。

## 任务
从简历文本/图片中提取结构化信息。

## 输出Schema
{basic_info: {name, phone, email, location, birth_year, work_years, target_position, self_summary},
 education: [{school, major, degree, start_date, end_date}],
 work_experience: [{company, position, duration, description}],
 projects: [{name, role, duration, description, achievements[], tech_stack[]}],
 skills: string[],
 ability_tags: {industry[], technology[], product[], capability[]}}''',
        acceptance=[
            '简历列表加载<2秒',
            'PDF/DOCX解析成功率≥95%',
            '版本对比正确高亮差异内容',
            '导出PDF格式排版美观',
        ],
        ui_desc='资产中心三栏布局：左栏为简历树形导航，中栏为简历预览/编辑器，右栏为AI优化建议。支持预览/编辑/对比三种模式切换。'
    )

    # F9
    add_feature_block(doc,
        fid='F9', name='AI增强简历编辑器', priority='P1', status_done=False,
        desc='模块化简历编辑器，每个section（教育、工作经历、项目、技能）独立编辑，支持hover显示AI优化建议（接受/重写/忽略），支持实时预览。',
        user_story='作为一名求职者，我希望在编辑简历时能获得实时AI建议，比如某段工作经历可以如何优化表述，帮我写得更专业。',
        preconditions=['已有简历数据'],
        main_flow=[
            '进入简历详情 > 编辑模式',
            '简历按模块展示：基本信息、教育背景、工作经历、项目经历、技能',
            '每个模块支持独立编辑（hover显示编辑按钮）',
            'hover任意段落，显示AI优化建议气泡',
            'AI建议包含：优化后文本、修改原因、预期效果',
            '用户可选择：接受（替换原文）、重写（AI生成新版本）、忽略',
            '底部显示结构化摘要：强化关键词数、新增量化表达数、待完善项数',
        ],
        alt_flows=['编辑过程中自动保存（每30秒或失焦时）'],
        exception_handling=[
            '异常E1：AI建议加载超时 → 显示"AI正在分析中..."占位',
            '异常E2：保存失败 → 本地缓存编辑内容 + 提示重试',
        ],
        output_spec='编辑后的完整简历JSON，包含各section的content和AI修改记录',
        acceptance=[
            'AI建议在hover后2秒内出现',
            '支持撤销/重做操作',
            '自动保存不丢失编辑内容',
        ],
        ui_desc='资产中心中栏编辑器。模块化卡片式设计，每个section为独立卡片，hover显示编辑和AI优化按钮。右侧显示AI建议面板：标题、原因、优先级、操作按钮。'
    )

    # F10
    add_feature_block(doc,
        fid='F10', name='STAR项目素材库', priority='P1', status_done=False,
        desc='使用STAR法则（Situation-Task-Action-Result）管理项目经历素材。每个项目包含情境、任务、行动步骤和成果数据，方便在不同简历和面试中复用。',
        user_story='作为一名求职者，我希望有一个地方用STAR法则整理我的项目经验，这样无论投哪个岗位，都能快速找到最合适的项目案例。',
        preconditions=['用户已登录'],
        main_flow=[
            '进入资产中心 > 项目素材库',
            '创建/编辑项目：标题、角色、时间段、STAR数据',
            'STAR数据：S(背景情境)、T(任务目标)、A(行动步骤列表)、R(量化成果)',
            '标注项目的技术栈和成就亮点',
            '支持排序和搜索',
            '在简历编辑和面试准备中可引用项目素材',
        ],
        alt_flows=['从简历中自动提取项目信息到素材库'],
        exception_handling=[
            '异常E1：STAR数据不完整 → 提示"建议补充Result部分的量化数据"',
        ],
        output_spec='{id, title, role, startDate, endDate, description, starData: {situation, task, action[], result}, techStack[], achievements[], sortOrder}',
        acceptance=[
            '支持至少20个项目管理',
            'STAR数据结构完整可编辑',
            '项目素材可在简历编辑器中一键引用',
        ],
        ui_desc='资产中心左侧导航 > 项目素材库。列表展示项目卡片（标题、角色、时间、技术标签）。点击进入STAR编辑界面。'
    )

    # F11
    add_feature_block(doc,
        fid='F11', name='成就证据库', priority='P2', status_done=False,
        desc='管理求职相关的证书、作品集和成就证据（证书照片、项目截图、数据报告等），方便在面试中提供佐证材料。',
        user_story='作为一名求职者，我希望能集中管理我的证书和作品，面试时能快速展示给面试官。',
        preconditions=['用户已登录'],
        main_flow=[
            '进入资产中心 > 成就证据库',
            '上传证书/作品：标题、描述、文件（图片/文档/链接）、分类、标签',
            '支持证书信息录入：证书名、发证机构、发证日期、有效期、证书编号',
            '作品集管理：封面图、描述、链接',
            '搜索和分类筛选',
        ],
        alt_flows=['支持批量上传多个文件'],
        exception_handling=[
            '异常E1：文件大小超限(>10MB) → 提示压缩后重新上传',
            '异常E2：文件格式不支持 → 提示支持的格式列表',
        ],
        output_spec='证书：{id, name, issuer, issueDate, expiryDate, certificateNo, fileUrl, category, level}; 作品：{id, title, description, fileUrl, fileType, coverUrl, tags[], viewCount}',
        acceptance=[
            '支持PNG/JPG/PDF/DOCX上传',
            '文件大小限制10MB',
            '列表支持网格/列表视图切换',
        ],
        ui_desc='资产中心左侧导航 > 成就证据库。网格式卡片展示证书/作品封面图，支持分类Tab切换（全部/证书/作品/其他）。'
    )


# ============================================================
# 主入口（Part 1）
# ============================================================

def main():
    print("[PRD V2.0] 开始生成文档...")
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    setup_styles(doc)
    add_header_footer(doc)

    print("  封面与目录...")
    add_cover(doc)
    add_revision_and_toc(doc)

    print("  第1章 产品综述...")
    add_chapter_1(doc)

    print("  第2章 系统架构...")
    add_chapter_2(doc)

    print("  第3章 机会工作台...")
    add_chapter_3(doc)

    print("  第4章 资产中心...")
    add_chapter_4(doc)

    # 保存中间文件
    output_path = '/home/user/webapp/webapp/prd_output/FindJob_PRD_V2.0_part1.docx'
    doc.save(output_path)
    print(f"  Part 1 保存成功: {output_path}")
    return doc


if __name__ == '__main__':
    main()
