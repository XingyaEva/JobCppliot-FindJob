#!/usr/bin/env python3
"""
Job Copilot 产品需求文档 (PRD) 生成器
参照 CareerAgent PRD 格式，基于项目实际代码分析生成
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ==================== 样式配置 ====================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 配置标题样式
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
    elif level == 2:
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
    elif level == 3:
        heading_style.font.size = Pt(12)
        heading_style.font.bold = True

def add_table_with_style(rows, cols, data, col_widths=None):
    """创建带样式的表格"""
    table = doc.add_table(rows=rows, cols=cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            # 设置字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            # 表头加粗
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                # 表头背景色
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50" w:val="clear"/>')
                cell._element.get_or_add_tcPr().append(shading)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    return table

def add_para(text, style_name='Normal', bold=False, color=None):
    """添加段落"""
    p = doc.add_paragraph(text, style=style_name)
    if bold:
        for run in p.runs:
            run.font.bold = True
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p

def add_bullet(text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

# ==================== 封面 ====================
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('产品需求文档 (PRD)')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph('')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Job Copilot - 智能求职助手 (V0.7.0)')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph('')
doc.add_paragraph('')

# 文档属性表
meta_data = [
    ['文档属性', '详情'],
    ['产品名称', 'Job Copilot - 智能求职助手'],
    ['版本号', 'V0.7.0 (Phase 0-9 已完成)'],
    ['状态', '已上线运营'],
    ['技术栈', 'Hono + TypeScript + Cloudflare Pages'],
    ['生产环境', 'https://job-copilot.pages.dev'],
    ['撰写人', 'AI Developer'],
    ['最后更新', '2026-02'],
    ['保密级别', '内部公开'],
]
add_table_with_style(len(meta_data), 2, meta_data, col_widths=[4, 12])

doc.add_page_break()

# ==================== 目录 ====================
doc.add_heading('目录 (Table of Contents)', level=1)
toc_items = [
    '1. 产品综述 (Product Overview)',
    '    1.1 背景与机会',
    '    1.2 产品定位与价值主张',
    '    1.3 核心用户画像',
    '    1.4 项目范围与版本规划',
    '2. 业务流程与架构 (Business Logic & Architecture)',
    '    2.1 总体业务流程图',
    '    2.2 Multi-Agent 架构设计',
    '    2.3 DAG 执行引擎',
    '    2.4 数据流向图',
    '3. 详细功能需求 (Detailed Functional Requirements)',
    '    3.1 [F1] JD 解析与深度透视',
    '    3.2 [F2] 简历解析与结构化',
    '    3.3 [F3] 人岗匹配度诊断报告',
    '    3.4 [F4] 面试准备与模拟演练',
    '    3.5 [F5] 简历定制化优化',
    '    3.6 [F6] JD 定向简历生成',
    '    3.7 [F7] 公司深度分析',
    '    3.8 [F8] 面试辅导与实时反馈',
    '    3.9 [F9] 投递跟踪管理',
    '    3.10 [F10] 模型评测仪表盘',
    '4. Prompt 工程策略 (Prompt Engineering Strategy)',
    '    4.1 System Prompt 设计体系',
    '    4.2 上下文管理 (Context Management)',
    '    4.3 幻觉控制与输出规范',
    '    4.4 成本优化策略',
    '5. 技术架构 (Technical Architecture)',
    '    5.1 前端架构',
    '    5.2 后端架构',
    '    5.3 LLM 双通道架构',
    '    5.4 PDF 智能解析路由',
    '    5.5 URL 爬虫模块',
    '6. 用户体验与交互设计 (UX/UI Design)',
    '    6.1 设计原则',
    '    6.2 响应式布局',
    '    6.3 交互亮点',
    '7. 非功能性需求 (Non-Functional Requirements)',
    '    7.1 性能要求',
    '    7.2 数据隐私与安全',
    '    7.3 兼容性',
    '    7.4 可用性与容错',
    '8. 数据埋点与考核指标 (Metrics)',
    '9. 已完成功能清单与版本记录',
    '10. 优化建议与下一步规划',
]
for item in toc_items:
    add_para(item)

doc.add_page_break()

# ==================== 第1章：产品综述 ====================
doc.add_heading('1. 产品综述 (Product Overview)', level=1)

doc.add_heading('1.1 背景与机会', level=2)
add_para('在当前就业市场中，"海投"已成为求职者的常态行为，招聘方普遍使用 ATS (Applicant Tracking Systems) 自动筛选简历。求职者面临三大核心困境：')
add_bullet('简历同质化：一份通用简历无法匹配不同 JD 的关键词和隐性需求，导致初筛通过率极低。')
add_bullet('JD 理解表面化：求职者往往只看到"显性要求"（技能、学历），忽略了 JD 背后的"潜台词"（岗位痛点、团队阶段、业务诉求）。')
add_bullet('面试准备碎片化：即使获得面试机会，缺乏系统性的基于岗位分析的面试准备策略和针对性问题预测。')
add_para('')
add_para('现有工具多为简单的"简历模板"或"通用 AI 润色"，缺乏基于 JD 语义深度理解的全链路求职辅助。LLM 的成熟使得以低成本实现"一对一职业咨询师"成为可能。')
add_para('')
add_para('Job Copilot 正是在这一背景下诞生的产品，它不是简单的简历美化工具，而是一个 Multi-Agent 协作的智能求职助手，覆盖从"岗位透视"到"面试通关"的完整闭环。')

doc.add_heading('1.2 产品定位与价值主张', level=2)
add_para('定位：基于 LLM Multi-Agent 的智能求职助手，面向 AI / 产品 / 技术类求职者，提供从"岗位解析"到"简历定制"到"面试通关"的全流程一站式辅助。', bold=True)
add_para('')
add_para('核心价值：')
add_bullet('透视 - 帮助用户看透 JD 背后的隐性需求、岗位痛点、团队阶段和核心 KPI。')
add_bullet('匹配 - 多维度人岗匹配诊断，精准识别优势、差距和风险点。')
add_bullet('定制 - 拒绝万金油简历，实现"一岗一历"，智能关键词植入与 STAR 扩写。')
add_bullet('演练 - 基于岗位特征的面试问题预测，含 PREP 结构化回答和行为指导。')
add_bullet('追踪 - 全流程投递跟踪管理，量化求职效率。')

doc.add_heading('1.3 核心用户画像', level=2)

persona_data = [
    ['角色', '特征', '痛点', '典型场景'],
    ['应届生 / 初级求职者',
     '经历少，不知道如何突出亮点；不熟悉 STAR 法则；面试准备不充分。',
     '"我的简历感觉很单薄"，"面试官提的问题我完全没准备"。',
     '只有实习/社团经验，想投互联网产品/技术岗，需要 AI 挖掘亮点并生成针对性回答。'],
    ['转行者 / 跨界求职者',
     '过往经验与目标岗位不匹配；简历关键词脱节；缺乏对目标行业的深入了解。',
     '"我的经验和 JD 要求差距太大"，"不知道怎么包装过往经历"。',
     '需要 AI 将"销售业绩"映射到"用户洞察能力"，并分析目标公司和行业特点。'],
    ['在职求职者 / 资深人士',
     '时间宝贵，无法针对每个 JD 改简历；面试技巧生疏；需要快速了解目标公司。',
     '"下班太累了不想从头改简历"，"很多年没面试了，不知道现在面试都问什么"。',
     '复制 JD，上传旧简历，一键生成定向简历版本，通勤路上看面试脚本。'],
]
add_table_with_style(len(persona_data), 4, persona_data, col_widths=[3, 4.5, 4.5, 4.5])

doc.add_heading('1.4 项目范围与版本规划', level=2)
add_para('当前版本 (V0.7.0) 已包含：', bold=True)
add_bullet('支持 JD 截图（OCR）、JD 文本粘贴、URL 爬取三种 JD 输入方式。')
add_bullet('PDF（数字型/扫描型 OCR）、Word、文本多格式简历上传与结构化解析。')
add_bullet('JD 双维度深度分析（A维度：速览 + B维度：深度拆解）。')
add_bullet('多维度人岗匹配诊断报告（含匹配等级、分数、维度雷达分析）。')
add_bullet('公司深度分析（公司画像、AI 场景分析、竞品分析、面试情报）。')
add_bullet('面试准备全套材料（自我介绍、项目推荐、分类面试题 + PREP 回答）。')
add_bullet('简历通用优化与 JD 定向简历生成。')
add_bullet('面试辅导实时反馈（逐题点评、评分、改进建议）。')
add_bullet('投递跟踪管理系统（状态流转、面试记录、数据统计）。')
add_bullet('模型评测仪表盘（Agent 性能、成本分析、A/B 实验框架）。')

add_para('')
add_para('不包含 / 计划中：')
add_bullet('多语言支持（当前仅中文输出）。')
add_bullet('真实语音模拟面试（仅文本）。')
add_bullet('职位推荐引擎。')
add_bullet('服务端持久化数据存储（当前使用 LocalStorage + 内存）。')
add_bullet('用户账号体系与多设备同步。')

doc.add_page_break()

# ==================== 第2章：业务流程与架构 ====================
doc.add_heading('2. 业务流程与架构 (Business Logic & Architecture)', level=1)

doc.add_heading('2.1 总体业务流程图', level=2)
add_para('Job Copilot 的业务流程被设计为线性但可回溯的结构，核心分为五层：')
add_para('')
add_para('输入层 (Input)', bold=True)
add_bullet('方式一：上传 JD 截图（支持粘贴板 Ctrl+V），系统自动 OCR 识别。')
add_bullet('方式二：粘贴 JD 文本，系统直接解析。')
add_bullet('方式三：粘贴招聘网站 URL（Boss直聘/拉勾/猎聘），系统自动爬取并解析。')
add_bullet('上传简历文件（PDF 数字型/扫描型、Word、文本）。')

add_para('')
add_para('解析层 (Processing)', bold=True)
add_bullet('JD 预处理 Agent：OCR 识别 / 文本清洗 → 结构化提取。')
add_bullet('JD 结构化 Agent：提取岗位名称、公司、薪资、职责、要求等结构化字段。')
add_bullet('简历预处理：智能路由（数字 PDF → pdfplumber 快速提取；扫描件 → MinerU OCR 深度解析）。')
add_bullet('简历结构化 Agent：LLM 提取基本信息、教育、工作经历、项目经验、技能标签。')

add_para('')
add_para('分析层 (Reasoning)', bold=True)
add_bullet('A维度分析 Agent（岗位速览）：技术栈密度、产品类型、业务领域、团队阶段。')
add_bullet('B维度分析 Agent（深度拆解）：行业背景、技术深度、产品经验、能力模型。')
add_bullet('匹配评估 Agent：多维度对比简历与 JD，生成匹配等级、分数和 Gap Analysis。')
add_bullet('公司分析 Agent：联网搜索，生成公司画像、AI 场景、竞品分析、面试情报。')

add_para('')
add_para('生成层 (Generation)', bold=True)
add_bullet('简历优化 Agent：根据 Gap 填补内容，关键词植入，STAR 扩写，动词升级。')
add_bullet('定向简历 Agent：基于 A/B 维度生成岗位定制版简历。')
add_bullet('面试准备 Agent：生成自我介绍、项目推荐、分类问题 + PREP 结构化回答。')
add_bullet('面试辅导 Agent：对用户答案进行实时评分、点评和改进建议。')

add_para('')
add_para('交付层 (Output)', bold=True)
add_bullet('在线查看完整解析报告、匹配诊断、面试材料。')
add_bullet('数据导出（JSON 格式完整数据包）。')
add_bullet('投递跟踪面板（状态管理、面试记录、统计报表）。')

doc.add_heading('2.2 Multi-Agent 架构设计', level=2)
add_para('Job Copilot 采用 Multi-Agent 协作模式，通过 DAG (有向无环图) 执行引擎调度多个垂直 Agent：')
add_para('')

agent_data = [
    ['Agent', '角色定位', '核心职责', '使用模型', '调用时机'],
    ['JD 预处理', '文本清洗师', 'OCR/文本 → 清洗后的 JD 文本', 'Qwen-VL-Max / Qwen-Plus', 'JD 输入时'],
    ['JD 结构化', '数据建模师', 'JD 文本 → 结构化 JSON', 'Qwen-Plus', 'JD 预处理后'],
    ['A维度分析', '速览分析师', '技术栈/产品类型/领域/阶段', 'Qwen-Plus', 'JD 结构化后'],
    ['B维度分析', '深度拆解师', '行业/技术/产品/能力四维度', 'Qwen-Plus', 'A维度分析后'],
    ['简历预处理', 'PDF 工程师', 'PDF/Word → 清洗后文本', 'pdfplumber / MinerU', '简历上传时'],
    ['简历结构化', '信息提取师', '文本 → 结构化简历 JSON', 'Qwen-Plus', '简历预处理后'],
    ['匹配评估', '诊断分析师', '简历 vs JD → 匹配等级/分数', 'Qwen-Plus', '用户触发匹配'],
    ['公司分析', '情报分析师', '公司画像/AI场景/竞品/面试情报', 'Qwen-Plus (联网)', '用户触发分析'],
    ['面试准备', '面试教练', '自我介绍/题目/PREP回答', 'Qwen-Plus', '用户触发准备'],
    ['简历优化', '优化顾问', '基于 Gap 重写简历', 'Qwen-Plus', '用户触发优化'],
    ['定向简历', '定制专家', '基于 A/B 维度生成定向简历', 'Qwen-Plus', '用户触发生成'],
    ['面试辅导', '陪练教练', '答案评分/点评/改进建议', 'Qwen-Plus', '用户回答后'],
]
add_table_with_style(len(agent_data), 5, agent_data, col_widths=[2.5, 2.2, 4, 2.5, 2.5])

doc.add_heading('2.3 DAG 执行引擎', level=2)
add_para('Job Copilot 使用自研的 DAG 执行引擎（dag-executor.ts）管理 Agent 依赖关系和并行执行：')
add_bullet('节点管理：每个 Agent 是一个 DAG 节点，声明依赖关系。')
add_bullet('并行执行：无依赖关系的节点自动并行调度。')
add_bullet('状态广播：实时通知前端各节点执行状态（pending / running / completed / error）。')
add_bullet('评测收集：自动收集每个 Agent 的执行指标（耗时、token 数、成本、成功率）。')
add_bullet('异步任务：支持 waitUntil 异步模式，避免 Cloudflare Workers 长连接超时。')
add_para('')
add_para('JD 解析 DAG 示例：')
add_para('  JD预处理 → JD结构化 → A维度分析 → B维度分析')
add_para('                            ↘ (并行) ↗')
add_para('')
add_para('简历解析 DAG 示例：')
add_para('  简历预处理 → 简历结构化解析')

doc.add_heading('2.4 数据流向图', level=2)
add_para('数据存储采用前端 LocalStorage + 后端内存缓存的混合模式：')
add_bullet('前端 LocalStorage：持久存储岗位列表、简历数据、匹配结果、面试记录、投递跟踪、评测指标。')
add_bullet('后端内存缓存：存储 DAG 执行状态、异步任务状态、解析进度。')
add_bullet('API 通信：前端通过 RESTful API 触发 Agent 执行，结果返回前端持久化。')
add_para('')
add_para('数据生命周期：')
add_para('  用户输入 → API 调用 → Agent 执行 → 结果返回前端 → LocalStorage 持久化 → 页面渲染')

doc.add_page_break()

# ==================== 第3章：详细功能需求 ====================
doc.add_heading('3. 详细功能需求 (Detailed Functional Requirements)', level=1)

# F1: JD解析
doc.add_heading('3.1 [F1] JD 解析与深度透视', level=2)
doc.add_heading('3.1.1 功能描述', level=3)
add_para('系统接受三种 JD 输入方式，通过 4 个 Agent 的 DAG 流水线，将非结构化的岗位描述转化为结构化分析报告。不仅提取"显性要素"（技能/学历/年限），更重要的是分析 JD 的"潜台词"——岗位痛点、团队阶段、核心 KPI 预测。')

doc.add_heading('3.1.2 输入方式', level=3)
add_bullet('截图上传：支持 JPG/PNG，支持拖拽、点击选择、Ctrl+V 粘贴板直接粘贴。')
add_bullet('文本粘贴：直接粘贴 JD 文本内容。')
add_bullet('URL 爬取（新）：粘贴 Boss直聘/拉勾/猎聘 URL，自动爬取并解析（支持 Cookie 配置）。')

doc.add_heading('3.1.3 DAG 处理流程', level=3)
add_para('节点1 - JD预处理：OCR 图片识别 / 文本清洗，去除噪声。')
add_para('节点2 - JD结构化：提取岗位名称、公司、薪资、地点、职责、要求、优先项等结构化字段。')
add_para('节点3 - A维度分析（岗位速览）：')
add_bullet('A1 技术栈密度：核心技术关键词、技术栈密度（高/中/低）、求职策略。')
add_bullet('A2 产品类型：B端/C端/G端/平台型等，及其对产品经理的差异化要求。')
add_bullet('A3 业务领域：主要业务领域及求职者需要储备的行业知识。')
add_bullet('A4 团队阶段：0-1探索期/1-10成长期/10-100成熟期，及对应的能力要求。')

add_para('节点4 - B维度分析（深度拆解）：')
add_bullet('B1 行业背景要求：是否必需、是否优先、具体行业、年限要求。')
add_bullet('B2 技术背景要求：学历门槛、技术深度分层（了解/熟悉/精通）。')
add_bullet('B3 产品经验要求：产品类型要求、全周期经验、0-1 经验。')
add_bullet('B4 能力模型要求：核心能力项及具体要求描述。')

doc.add_heading('3.1.4 交互设计', level=3)
add_bullet('PC端三栏布局：左栏（原图/原文本+岗位链接）+ 中栏（结构化分析结果）+ 右栏（解析文字+关键词高亮）。')
add_bullet('移动端单栏布局：基本信息 → A维度 → B维度 → 原始JD（折叠）。')
add_bullet('实时 DAG 进度显示：四个节点的状态（等待/执行中/完成）实时更新。')
add_bullet('操作按钮栏：投递、匹配分析、定向简历、面试准备、简历优化。')

doc.add_heading('3.1.5 异常处理', level=3)
add_bullet('图片 OCR 失败 → 提示用户使用文本粘贴模式。')
add_bullet('URL 爬取失败 → 返回具体错误原因（平台不支持/Cookie 过期/网络问题），引导手动输入。')
add_bullet('解析超时（3分钟） → 提示超时并建议重试。')

# F2: 简历解析
doc.add_heading('3.2 [F2] 简历解析与结构化', level=2)
doc.add_heading('3.2.1 功能描述', level=3)
add_para('系统采用智能 PDF 路由策略，将用户上传的简历文档转化为结构化 JSON 数据对象，涵盖基本信息、教育背景、工作经历、项目经验、技能标签等维度。')

doc.add_heading('3.2.2 智能解析路由', level=3)
add_para('Phase 3 引入的智能路由策略：')
add_bullet('数字 PDF（约 90% 简历）→ Python pdfplumber 快速提取（5-10秒）。')
add_bullet('扫描件 PDF → MinerU VLM + OCR 深度解析（45-60秒）。')
add_bullet('前端 pdf.js 提取（备选）→ 无需后端服务，秒级响应。')
add_bullet('文本直接输入 → 直接进入 LLM 结构化。')

doc.add_heading('3.2.3 结构化输出', level=3)
add_bullet('基本信息：姓名、联系方式、求职意向。')
add_bullet('教育背景：学校、专业、学历、时间。')
add_bullet('工作经历：公司、职位、时间、职责描述。')
add_bullet('项目经验：项目名称、角色、技术栈、成果。')
add_bullet('技能标签：技术技能、工具、语言等。')
add_bullet('能力标签：AI 推断的核心能力标签。')

doc.add_heading('3.2.4 版本管理', level=3)
add_para('简历支持完整的版本管理体系：')
add_bullet('基础版（v1）：原始上传解析的简历。')
add_bullet('优化版：通用 AI 优化后的版本。')
add_bullet('定向版：基于特定 JD 生成的定制版本，标注关联岗位。')
add_bullet('手动编辑版：用户手动修改后自动创建版本快照。')

# F3: 匹配诊断
doc.add_heading('3.3 [F3] 人岗匹配度诊断报告', level=2)
doc.add_heading('3.3.1 功能描述', level=3)
add_para('在优化简历前，先给用户精准的匹配度诊断，明确"你在哪里"和"差距在哪里"。')

doc.add_heading('3.3.2 输出内容', level=3)
add_bullet('匹配等级：S/A/B/C/D 五级评估。')
add_bullet('匹配分数：百分制精确评分。')
add_bullet('维度匹配：多维度详细匹配（专业技能/行业经验/学历背景/产品能力/软素质）。')
add_bullet('优势列表：简历中完全匹配 JD 要求的亮点。')
add_bullet('差距分析：JD 要求但简历未体现的关键能力。')
add_bullet('面试聚焦建议：基于 Gap 给出面试重点准备方向。')

doc.add_heading('3.3.3 匹配等级规则', level=3)
match_level_data = [
    ['等级', '分数范围', '含义', '建议行动'],
    ['S', '90-100', '完美匹配', '直接投递，重点准备面试'],
    ['A', '75-89', '高度匹配', '微调简历后投递'],
    ['B', '60-74', '中度匹配', '需要针对性优化简历'],
    ['C', '40-59', '低度匹配', '建议深度改写简历或考虑其他岗位'],
    ['D', '0-39', '不匹配', '建议重新评估目标岗位'],
]
add_table_with_style(len(match_level_data), 4, match_level_data, col_widths=[2, 2.5, 3, 6])

# F4: 面试准备
doc.add_heading('3.4 [F4] 面试准备与模拟演练', level=2)
doc.add_heading('3.4.1 功能描述', level=3)
add_para('基于岗位分析、简历解析和匹配评估的综合结果，生成全面的面试准备材料。不是泛泛而谈，而是高度针对性的"作战手册"。')

doc.add_heading('3.4.2 输出结构', level=3)
add_para('一、自我介绍（三个版本）：')
add_bullet('30秒版：电梯演讲，核心亮点。')
add_bullet('1分钟版：标准面试自我介绍。')
add_bullet('3分钟版：详细版，含项目亮点。')

add_para('')
add_para('二、项目推荐：')
add_bullet('从简历中推荐最匹配目标岗位的 2-3 个项目。')
add_bullet('每个项目标注推荐理由和面试讲述要点。')

add_para('')
add_para('三、分类面试题 + PREP 结构化回答：')
add_bullet('技术能力题：基于 JD 技术栈要求。')
add_bullet('产品/业务题：基于岗位职责和业务领域。')
add_bullet('行为面试题：基于简历项目经历。')
add_bullet('压力/情景题：基于岗位特点。')
add_para('每道题包含：问题原文、考察意图、PREP 结构化回答（Point → Reason → Example → Point）。')

add_para('')
add_para('四、整体面试策略：')
add_bullet('核心展示主题。')
add_bullet('关键差异化优势。')
add_bullet('需要主动引导的话题。')
add_bullet('需要回避或巧妙转化的弱点。')

# F5: 简历优化
doc.add_heading('3.5 [F5] 简历定制化优化', level=2)
doc.add_heading('3.5.1 功能描述', level=3)
add_para('根据匹配诊断的 Gap Analysis 结果，智能优化用户简历内容。这是产品的核心功能之一。')

doc.add_heading('3.5.2 优化策略', level=3)
add_bullet('关键词植入：将 JD 中的高频关键词自然融入简历描述中（非机械堆砌）。')
add_bullet('STAR 扩写：检查每段经历，补充 Situation/Task/Action/Result 缺失部分。')
add_bullet('动词升级：将"做了"、"负责"替换为"主导"、"构建"、"驱动"、"优化"等强动词。')
add_bullet('成果量化：引导量化工作成果（提升了 XX%，节约了 XX 万）。')
add_bullet('去噪处理：删除与目标岗位无关的冗余信息。')
add_bullet('用户自定义建议：支持用户输入额外优化方向。')

doc.add_heading('3.5.3 输出格式', level=3)
add_bullet('优化后的完整简历各模块（基本信息、教育、工作经历、项目、技能）。')
add_bullet('详细的修改说明（每处修改标注原因和策略）。')
add_bullet('匹配度预估提升百分比。')

# F6: 定向简历
doc.add_heading('3.6 [F6] JD 定向简历生成', level=2)
add_para('Phase 7 新增功能。与 F5 通用优化不同，定向简历是完全基于特定 JD 的 A/B 维度分析结果，生成高度针对性的简历版本。')
add_para('')
add_para('优化维度：')
add_bullet('关键词强化：基于 A1 技术栈密度注入关键词。')
add_bullet('经历排序：基于 B 维度匹配度重新排列经历优先级。')
add_bullet('内容调整：基于 A4 团队阶段和 B4 能力模型调整描述重点。')
add_bullet('弱化/删除建议：识别与目标岗位无关的内容。')
add_para('')
add_para('输出：新的简历版本（自动关联到目标岗位），包含版本标签和预估匹配度提升。')

# F7: 公司分析
doc.add_heading('3.7 [F7] 公司深度分析', level=2)
add_para('联网搜索目标公司信息，生成全面的公司画像报告，帮助求职者在面试前深入了解目标公司。')
add_para('')
add_para('输出内容：')
add_bullet('公司画像：基本信息、发展阶段、融资情况、企业文化。')
add_bullet('AI 场景分析：公司在 AI 领域的布局、技术栈、业务应用。')
add_bullet('竞品分析：同赛道竞争对手对比。')
add_bullet('面试情报：面试流程、常见问题方向、价值观关键词。')
add_bullet('一句话总结：精练概括公司定位和特点。')

# F8: 面试辅导
doc.add_heading('3.8 [F8] 面试辅导与实时反馈', level=2)
add_para('用户可以针对面试题进行回答练习，AI 会实时评价并给出改进建议。')
add_para('')
add_para('功能特点：')
add_bullet('逐题回答：用户可以逐题回答面试准备中的问题。')
add_bullet('即时评分：1-10 分评分，含详细评语。')
add_bullet('must_fix 高优先级建议：必须修正的关键问题。')
add_bullet('改进方向：整体改进方向建议。')
add_bullet('优化后参考回答：AI 生成的改进版本。')
add_bullet('亮点标注：回答中的闪光点肯定。')
add_bullet('批量辅导：支持多题并行评价。')

# F9: 投递跟踪
doc.add_heading('3.9 [F9] 投递跟踪管理', level=2)
add_para('Phase 9 新增功能。完整的求职流程管理工具。')
add_para('')
add_para('核心功能：')
add_bullet('状态流转：已投递 → 简历通过 → 一面 → 二面 → HR面 → Offer → 入职（支持自定义）。')
add_bullet('面试记录：记录每轮面试的时间、面试官、问题回顾、自我评价。')
add_bullet('数据看板：投递总数、各状态分布、面试率、Offer 率等关键指标。')
add_bullet('快捷投递：岗位详情页一键创建投递记录。')

# F10: 评测仪表盘
doc.add_heading('3.10 [F10] 模型评测仪表盘', level=2)
add_para('Phase 6 新增功能。面向开发者和产品运营的 Agent 性能监控面板。')
add_para('')
add_para('监控维度：')
add_bullet('Agent 性能：每个 Agent 的调用次数、成功率、平均耗时。')
add_bullet('模型成本：各模型的 Token 消耗和估算费用。')
add_bullet('A/B 实验：支持配置实验分组，对比不同模型/Prompt 的效果。')
add_bullet('模型定价：内置 GPT-4.1/4o/4o-mini、Qwen-Max/Plus/Turbo、DeepSeek-V3 等模型单价表。')

doc.add_page_break()

# ==================== 第4章：Prompt工程 ====================
doc.add_heading('4. Prompt 工程策略 (Prompt Engineering Strategy)', level=1)

doc.add_heading('4.1 System Prompt 设计体系', level=2)
add_para('所有 Agent 的 System Prompt 采用统一的三模块架构：')
add_para('')
add_para('模块一：角色定义 (Role Definition)', bold=True)
add_para('为每个 Agent 设定清晰的专家人设，例如：')
add_bullet('匹配评估 Agent："你是一位拥有 10 年以上经验的资深 HR 总监和人才评估专家。"')
add_bullet('面试准备 Agent："你是一位资深面试教练，精通 STAR 和 PREP 方法论。"')
add_bullet('简历优化 Agent："你是一位拥有 15 年经验的资深简历优化专家和 HR 顾问。"')

add_para('')
add_para('模块二：任务约束 (Task Constraints)', bold=True)
add_bullet('输出格式：严格 JSON，不含 Markdown 标记。')
add_bullet('语言要求：中文输出，技术名词保留英文。')
add_bullet('长度控制：summary ≤ 50 字，description ≤ 100 字，列表 ≤ 5 项。')
add_bullet('禁止编造：不得编造用户没有做过的公司、职位或项目。缺失数据使用占位符。')
add_bullet('缺失标记：信息不足时标注"信息不足"，不捏造。')

add_para('')
add_para('模块三：知识库注入 (Knowledge Base)', bold=True)
add_bullet('STAR 法则定义和示例。')
add_bullet('PREP 结构化回答框架。')
add_bullet('匹配等级评判标准（S/A/B/C/D）。')
add_bullet('行业特征知识（AI/SaaS/电商等领域特点）。')

doc.add_heading('4.2 上下文管理 (Context Management)', level=2)
add_para('采用 Map-Reduce 策略管理长文本上下文：')
add_bullet('Map 阶段：独立总结 JD（JD Profile）和简历（Resume Profile），各自压缩为精简结构化数据。')
add_bullet('Reduce 阶段：将两个 Profile 作为输入传给下游 Agent（匹配/优化/面试），避免 Token 溢出。')
add_bullet('选择性输入：面试准备仅提取"项目经历"部分作为 Context，减少噪音。')
add_bullet('DAG 依赖传递：上游 Agent 的输出自动作为下游 Agent 的输入。')

doc.add_heading('4.3 幻觉控制与输出规范', level=2)
add_para('为保证 AI 输出的可靠性，采用多重控制措施：')
add_bullet('结构化输出约束：所有 Agent 使用 JSON Mode，通过 Schema 定义严格约束输出格式。')
add_bullet('字段校验与默认值：每个 Agent 输出后进行字段存在性校验，缺失字段填充合理默认值。')
add_bullet('Token 优化指令：Prompt 中明确限定各字段的最大长度和列表项数。')
add_bullet('缺失信息标记：明确要求"信息不足"场景的处理方式，不得编造事实性数据。')
add_bullet('UI 免责提示：前端关键位置标注"AI 生成内容可能包含推测，请务必核实"。')

doc.add_heading('4.4 成本优化策略', level=2)
add_para('Phase 6 实现的 Prompt 和成本优化：')
add_bullet('Token 估算：内置 4 字符/token 的估算模型，实时计算 Agent 成本。')
add_bullet('模型选择策略：简单任务用 Qwen-Turbo（低成本），复杂任务用 Qwen-Plus（高质量）。')
add_bullet('输出长度约束：通过 Prompt 中的 Token 优化指令控制输出大小。')
add_bullet('A/B 实验：通过评测仪表盘对比不同模型在质量/成本间的平衡。')

model_cost_data = [
    ['模型', '输入价格 ($/1M tokens)', '输出价格 ($/1M tokens)', '适用场景'],
    ['gpt-4.1', '$2.00', '$8.00', '高质量分析（备用通道）'],
    ['gpt-4o', '$2.50', '$10.00', '复杂推理（备用通道）'],
    ['gpt-4o-mini', '$0.15', '$0.60', '轻量任务（备用通道）'],
    ['qwen-max', '$20.00', '$60.00', '最高质量（保留未用）'],
    ['qwen-plus', '$3.00', '$6.00', '主力模型（大部分 Agent）'],
    ['qwen-turbo', '$3.00', '$6.00', '轻量/快速任务'],
    ['deepseek-v3', '$0.27', '$1.10', '高性价比（VectorEngine通道）'],
]
add_table_with_style(len(model_cost_data), 4, model_cost_data, col_widths=[3, 4, 4, 5.5])

doc.add_page_break()

# ==================== 第5章：技术架构 ====================
doc.add_heading('5. 技术架构 (Technical Architecture)', level=1)

doc.add_heading('5.1 前端架构', level=2)
add_bullet('框架：Hono JSX（服务端渲染 SSR）。')
add_bullet('样式：TailwindCSS CDN。')
add_bullet('图标：FontAwesome 6.4。')
add_bullet('数据持久化：浏览器 LocalStorage（带命名空间隔离）。')
add_bullet('交互：原生 JavaScript DOM 操作（无前端框架依赖）。')
add_bullet('响应式：PC 三栏布局 + 移动端单栏自适应。')

doc.add_heading('5.2 后端架构', level=2)
add_bullet('框架：Hono v4 (TypeScript)。')
add_bullet('运行时：Cloudflare Workers（边缘计算）。')
add_bullet('构建：Vite + @hono/vite-cloudflare-pages 插件。')
add_bullet('部署：Cloudflare Pages（全球 CDN + 边缘执行）。')
add_bullet('API 风格：RESTful，JSON 数据交换。')

doc.add_heading('5.3 LLM 双通道架构', level=2)
add_para('系统采用双 LLM 通道设计，确保可用性和灵活性：')
add_para('')
channel_data = [
    ['通道', 'API Provider', 'Base URL', '特点'],
    ['主通道', '阿里云百炼 (DashScope)', 'dashscope.aliyuncs.com', '低成本、国内稳定、Qwen 模型族'],
    ['备用通道', 'VectorEngine', 'api.vectorengine.ai', 'GPT-4 系列、DeepSeek，国际模型'],
]
add_table_with_style(len(channel_data), 4, channel_data, col_widths=[2.5, 4, 5, 5])
add_para('')
add_para('每个 Agent 可独立配置：provider（通道选择）、model（模型选择）、temperature、maxTokens、jsonMode、enableSearch（联网搜索）。')

doc.add_heading('5.4 PDF 智能解析路由', level=2)
add_para('Phase 3 设计的三层 PDF 解析策略：')
add_para('')
pdf_data = [
    ['层级', '工具', '适用场景', '速度', '准确率'],
    ['快速层', 'pdfplumber (Python)', '数字型 PDF (~90%)', '5-10 秒', '高'],
    ['深度层', 'MinerU VLM + OCR', '扫描件 PDF (~10%)', '45-60 秒', '极高'],
    ['前端层', 'pdf.js (浏览器)', '降级方案 / 即时预览', '<1 秒', '中'],
]
add_table_with_style(len(pdf_data), 5, pdf_data, col_widths=[2, 3.5, 4, 2.5, 2])
add_para('')
add_para('路由逻辑：先检测 Python 服务健康 → 快速分析 PDF 类型（数字/扫描）→ 路由到对应解析器。')

doc.add_heading('5.5 URL 爬虫模块', level=2)
add_para('支持从招聘网站自动爬取 JD 内容：')
add_bullet('支持平台：Boss直聘、拉勾、猎聘。')
add_bullet('爬虫方案：ScraperAPI（主选）/ ScrapingBee（备选）。')
add_bullet('CSS 选择器：每个平台预配置标题、公司、JD内容、薪资、地点等字段的提取规则。')
add_bullet('Cookie 管理：支持用户配置平台登录 Cookie，用于访问需要认证的页面。')
add_bullet('缺失字段检测：自动检测爬取数据完整性，缺失时给出明确警告和建议。')

doc.add_page_break()

# ==================== 第6章：UX/UI 设计 ====================
doc.add_heading('6. 用户体验与交互设计 (UX/UI Design)', level=1)

doc.add_heading('6.1 设计原则', level=2)
add_bullet('极简主义：黑白灰为主色调，减少视觉噪音。')
add_bullet('渐进式披露：复杂信息（如 B 维度分析）采用折叠展开交互。')
add_bullet('即时反馈：所有 AI 操作都有实时进度展示（DAG 节点状态、百分比进度条）。')
add_bullet('容错友好：错误提示清晰，提供明确的下一步操作建议。')

doc.add_heading('6.2 响应式布局', level=2)
add_bullet('PC端（≥1024px）：三栏布局，充分利用大屏空间。')
add_bullet('平板端（768-1024px）：双栏布局，侧栏可折叠。')
add_bullet('移动端（<768px）：单栏堆叠布局，Tab 切换功能区域。')

doc.add_heading('6.3 交互亮点', level=2)
add_bullet('剪贴板粘贴截图：Ctrl+V 直接粘贴 JD 截图，减少操作步骤。')
add_bullet('拖拽上传：支持拖拽文件到上传区域。')
add_bullet('关键词高亮：JD 解析后在原文中高亮标注识别出的技术关键词。')
add_bullet('骨架屏加载：数据加载时展示骨架屏动画，减少感知等待时间。')
add_bullet('Toast 通知：操作完成后的轻量级通知提示。')
add_bullet('步骤引导：首页五步流程指引（解析岗位→上传简历→匹配评估→面试准备→优化简历），已完成步骤自动标绿。')
add_bullet('Tab 切换：新建岗位页支持"手动输入"和"URL爬取"两种模式无缝切换。')

doc.add_page_break()

# ==================== 第7章：非功能性需求 ====================
doc.add_heading('7. 非功能性需求 (Non-Functional Requirements)', level=1)

doc.add_heading('7.1 性能要求', level=2)
perf_data = [
    ['性能指标', '目标值', '当前实际', '优化措施'],
    ['JD 文本解析', '< 15 秒（端到端）', '约 10-15 秒', '4节点 DAG 并行 + 异步任务'],
    ['JD 截图 OCR 解析', '< 20 秒', '约 15-20 秒', 'Qwen-VL-Max 视觉模型'],
    ['简历解析（数字PDF）', '< 15 秒', '约 8-12 秒', 'pdfplumber 快速提取 + LLM 结构化'],
    ['简历解析（扫描件）', '< 90 秒', '约 50-70 秒', 'MinerU VLM + OCR + LLM'],
    ['匹配评估', '< 15 秒', '约 8-12 秒', '单 Agent 调用'],
    ['面试准备', '< 30 秒', '约 15-25 秒', '长输出任务，Token 优化'],
    ['简历优化', '< 20 秒', '约 10-15 秒', '单 Agent 调用'],
    ['感知速度', '< 0.5 秒', '< 0.1 秒', 'Phase 2 异步模式 + 即时进度反馈'],
]
add_table_with_style(len(perf_data), 4, perf_data, col_widths=[4, 3, 3, 6])

doc.add_heading('7.2 数据隐私与安全', level=2)
add_bullet('前端存储：所有用户数据存储在浏览器 LocalStorage，不上传到任何第三方存储。')
add_bullet('服务端无状态：后端不持久存储用户简历和岗位数据（仅内存缓存用于处理流程）。')
add_bullet('文件处理：上传的简历文件在解析完成后不保留原文件，仅保留结构化结果。')
add_bullet('API Key 保护：LLM API Key 仅存于服务端环境变量，前端无法访问。')
add_bullet('数据导出/清空：用户可随时导出全量数据或一键清空所有数据。')

doc.add_heading('7.3 兼容性', level=2)
add_bullet('简历格式：PDF（数字型 + 扫描型 OCR）、Word(.docx)、纯文本、Markdown。')
add_bullet('JD 输入：截图（JPG/PNG）、文本、URL（Boss直聘/拉勾/猎聘）。')
add_bullet('浏览器：Chrome 90+、Firefox 88+、Safari 14+、Edge 90+（推荐 Chrome）。')
add_bullet('移动端：支持移动端浏览器查看报告和面试材料。')

doc.add_heading('7.4 可用性与容错', level=2)
add_bullet('LLM 双通道容灾：主通道（阿里云百炼）故障时自动切换到备用通道（VectorEngine）。')
add_bullet('PDF 解析降级：Python 服务不可用 → 自动降级到 MinerU → 降级到前端 pdf.js。')
add_bullet('URL 爬虫容错：平台不支持时引导手动输入；Cookie 过期时提示重新配置。')
add_bullet('异步任务超时：最多轮询 3 分钟（180次），超时后提示用户重试。')
add_bullet('Cloudflare Workers 兼容：使用 waitUntil 处理长任务，避免连接超时。')

doc.add_page_break()

# ==================== 第8章：数据指标 ====================
doc.add_heading('8. 数据埋点与考核指标 (Metrics)', level=1)
add_para('Job Copilot 通过内置的评测框架收集 Agent 运行指标，用于衡量产品效果和优化方向。')
add_para('')

metrics_data = [
    ['指标类型', '指标名称', '定义与计算方式', '业务含义'],
    ['北极星指标', '完整优化率', '(完成"导出数据"或"生成面试题"的用户) / (上传简历的用户)', '衡量产品是否跑通了用户价值闭环。'],
    ['过程指标', 'Agent 成功率', '各 Agent 调用成功次数 / 总调用次数', '监控 AI 管道可靠性。'],
    ['过程指标', '平均解析耗时', '各 Agent 的平均执行时间（毫秒）', '监控性能退化。'],
    ['成本指标', '单次调用成本', '基于 Token 消耗和模型定价估算', '优化模型选择的 ROI。'],
    ['活跃指标', '面试题回访率', '生成面试题后 24h 内再次访问该页面的比例', '说明用户真的在用材料备战。'],
    ['满意度', '功能覆盖率', '用户使用的功能数 / 总功能数', '衡量功能发现率。'],
    ['实验指标', 'A/B 组效果对比', '实验组 vs 对照组的成功率/质量差异', '验证 Prompt/模型优化效果。'],
]
add_table_with_style(len(metrics_data), 4, metrics_data, col_widths=[2.5, 3, 5, 5])

doc.add_page_break()

# ==================== 第9章：已完成功能 ====================
doc.add_heading('9. 已完成功能清单与版本记录', level=1)

phase_data = [
    ['阶段', '名称', '核心功能', '状态'],
    ['Phase 0', '框架搭建', 'Hono + Cloudflare Pages 基础架构、路由设计、Agent 框架', '已完成'],
    ['Phase 1', 'JD 解析管道', 'JD 预处理、结构化、A/B 维度分析 DAG', '已完成'],
    ['Phase 2', '简历解析', '简历上传、MinerU 集成、结构化提取、实时进度', '已完成'],
    ['Phase 3', '匹配评估', '多维度人岗匹配、匹配等级/分数、Gap Analysis', '已完成'],
    ['Phase 4', '面试准备', '自我介绍、项目推荐、分类面试题 + PREP 回答', '已完成'],
    ['Phase 5', '简历优化', '基于 Gap 的智能优化、关键词植入、STAR 扩写', '已完成'],
    ['Phase 6', '评测框架', '模型评测仪表盘、A/B 实验、成本分析、Prompt 优化', '已完成'],
    ['Phase 7', '定向简历', 'JD 定向简历生成、版本管理、简历关联岗位', '已完成'],
    ['Phase 8', '公司分析', '公司画像、AI 场景分析、竞品分析、面试情报', '已完成'],
    ['Phase 8.5', '面试辅导', '逐题辅导、答案评分、改进建议、批量辅导', '已完成'],
    ['Phase 9', '投递跟踪', '状态管理、面试记录、数据看板、快捷投递', '已完成'],
    ['P1 优化', 'MinerU 优化', 'VLM 模型、关闭冗余 OCR、参数优化', '已完成'],
    ['P2 优化', '用户体验', '异步模式（感知速度 80s→0.1s）、进度条、桌面通知', '已完成'],
    ['URL 爬虫', 'JD URL 爬取', 'Boss直聘/拉勾/猎聘 URL 自动爬取、Cookie 管理', '已完成'],
    ['统一导航', 'UI 重构', '全站统一导航栏、数据导出/清空/删除', '已完成'],
]
add_table_with_style(len(phase_data), 4, phase_data, col_widths=[2.5, 3, 7, 2])

doc.add_page_break()

# ==================== 第10章：优化建议 ====================
doc.add_heading('10. 优化建议与下一步规划', level=1)

doc.add_heading('10.1 架构升级建议（高优先级）', level=2)

add_para('10.1.1 服务端持久化存储', bold=True)
add_para('当前问题：所有用户数据存储在浏览器 LocalStorage，存在数据丢失风险（清除浏览器数据）、无法跨设备同步、无法做全局数据分析。')
add_para('建议方案：引入 Cloudflare D1 数据库（SQLite）存储核心业务数据（岗位、简历、匹配结果），配合 KV 存储会话缓存。预计工作量：2-3 周。')

add_para('')
add_para('10.1.2 用户账号体系', bold=True)
add_para('当前问题：无用户登录，数据绑定到单一浏览器。')
add_para('建议方案：接入 Cloudflare Access 或第三方 OAuth（微信/GitHub），实现用户注册登录和数据隔离。预计工作量：1-2 周。')

add_para('')
add_para('10.1.3 流式输出优化', bold=True)
add_para('当前问题：所有 LLM 响应为阻塞式等待，用户需等待完整结果。')
add_para('建议方案：对面试准备、简历优化等长输出任务实现 SSE (Server-Sent Events) 流式输出，逐段展示结果。预计工作量：1 周。')

doc.add_heading('10.2 功能增强建议（中优先级）', level=2)

add_para('10.2.1 简历编辑器', bold=True)
add_para('当前问题：用户无法在线编辑 AI 优化后的简历内容。')
add_para('建议方案：内嵌可视化简历编辑器，支持拖拽排序、内联编辑、实时预览和 PDF/Word 导出。')

add_para('')
add_para('10.2.2 多轮面试对话模拟', bold=True)
add_para('当前问题：面试辅导是单轮评价模式。')
add_para('建议方案：实现多轮面试对话，AI 扮演面试官，根据用户回答动态追问，模拟真实面试场景。')

add_para('')
add_para('10.2.3 更多招聘平台支持', bold=True)
add_para('当前问题：URL 爬取仅支持三个平台。')
add_para('建议方案：扩展支持智联招聘、前程无忧、LinkedIn 等平台，并优化 CSS 选择器鲁棒性。')

add_para('')
add_para('10.2.4 简历导出功能', bold=True)
add_para('当前问题：优化后的简历仅以结构化数据形式存在，无法导出为 PDF/Word。')
add_para('建议方案：基于 Puppeteer/jsPDF 实现简历排版和 PDF 生成功能。')

doc.add_heading('10.3 技术优化建议（持续迭代）', level=2)
add_bullet('Prompt 持续优化：基于评测仪表盘数据，迭代优化各 Agent Prompt 的输出质量。')
add_bullet('模型升级评估：持续跟踪 Qwen3/GPT-5/Claude-4 等新模型发布，评估替换收益。')
add_bullet('缓存策略：对相同 JD 的分析结果做缓存，避免重复 LLM 调用。')
add_bullet('监控告警：接入 Cloudflare Analytics，监控 API 错误率、延迟、Agent 异常。')
add_bullet('国际化：为未来多语言支持预埋 i18n 架构。')

doc.add_heading('10.4 下一阶段路线图 (Roadmap)', level=2)

roadmap_data = [
    ['阶段', '时间', '核心目标', '关键交付'],
    ['V0.8', '2026 Q1', '数据持久化 + 账号体系', 'Cloudflare D1 迁移、OAuth 登录、跨设备同步'],
    ['V0.9', '2026 Q2', '内容创作增强', '简历编辑器、PDF 导出、模板库'],
    ['V1.0', '2026 Q3', '正式发布版', '多轮面试模拟、更多平台支持、性能全面优化'],
    ['V1.1', '2026 Q4', '增长与变现', '用户增长策略、付费功能、社区运营'],
]
add_table_with_style(len(roadmap_data), 4, roadmap_data, col_widths=[2, 2.5, 4, 7.5])

doc.add_paragraph('')
doc.add_paragraph('')

# 文档结尾
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = end.add_run('--- 文档结束 ---')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

end2 = doc.add_paragraph()
end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = end2.add_run(f'生成时间：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# 保存文件
output_path = '/home/user/webapp/Job_Copilot_PRD_V0.7.0.docx'
doc.save(output_path)
print(f'PRD 文档已生成: {output_path}')
print(f'文件大小: {__import__("os").path.getsize(output_path) / 1024:.1f} KB')
